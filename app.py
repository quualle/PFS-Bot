import os
import json
import re
import time
import logging
import traceback
from datetime import datetime, timedelta, date
import calendar # Import für Monatsberechnungen
from functools import wraps
import uuid
import tempfile
import requests  # Added import for requests
from datetime import datetime, timedelta
import dateparser
from conversation_manager import ConversationManager

from flask import (
    Flask, render_template, request, redirect, url_for, flash, jsonify, session, Response
)
from flask_wtf import CSRFProtect
from flask_session import Session
from dotenv import load_dotenv
from google.cloud import storage
from google.oauth2 import service_account
from google.cloud import bigquery
from werkzeug.utils import secure_filename
import openai
import tiktoken
import yaml
import re
from sql_query_helper import apply_query_enhancements
from query_router import determine_query_approach, determine_function_need, handle_conversational_clarification, process_user_query
from extract import (format_customer_details, format_date, extract_date_params, 
                    extract_enhanced_date_params, extract_agency_name, 
                    extract_customer_name, extract_parameters_with_llm, 
                    extract_enhanced_parameters)
from bigquery_functions import (
    get_user_id_from_email, 
    get_bigquery_client, 
    get_leads_for_seller, 
    get_contracts_for_seller, 
    get_households_for_seller, 
    calculate_kpis_for_seller,
    get_seller_data,
    execute_bigquery_query,
    format_query_result)
from tool_manager import load_tool_config, create_tool_description_prompt, select_tool, load_tool_descriptions, load_tool_descriptions, select_optimal_tool_with_reasoning
try:
    from query_selector import select_query_with_llm, update_selection_feedback, process_clarification_response, process_text_clarification_response
    USE_LLM_QUERY_SELECTOR = True
    logging.info("LLM-based query selector loaded successfully")
except ImportError as e:
    logging.warning(f"LLM-based query selector not available: {e}")
    USE_LLM_QUERY_SELECTOR = False
from llm_manager import create_enhanced_system_prompt, generate_fallback_response, call_llm
from utils import debug_print

def load_tool_config():
    """Liefert die Standard-Tool-Konfiguration"""
    # Direkte Rückgabe der Standardkonfiguration ohne Datei-Zugriff
    return {
        "description": "Tool-Konfiguration für LLM-basierte Entscheidungen",
        "fallback_tool": "get_care_stays_by_date_range",
        "use_llm_selection": True
    }

conversation_manager = ConversationManager(max_history=10)


###########################################
# PINECONE: gRPC-Variante laut Quickstart
###########################################
from pinecone.grpc import PineconeGRPC as Pinecone
from pinecone import ServerlessSpec

# Zusätzliche Importe für Dateiverarbeitung
from PyPDF2 import PdfReader
import docx

# Für Datum / Statistik
from datetime import datetime

from bigquery_functions import (
    handle_function_call, 
    summarize_query_result, 
    get_user_id_from_email
)

# Laden der Umgebungsvariablen aus .env
load_dotenv()

# Google OAuth Konfiguration
def configure_google_auth(app):
    """Konfiguriert die direkte Google OAuth-Authentifizierung."""
    # Stellen Sie sicher, dass die Umgebungsvariablen gesetzt sind
    if not os.getenv('GOOGLE_CLIENT_ID') or not os.getenv('GOOGLE_CLIENT_SECRET'):
        print("WARNUNG: Google OAuth-Credentials nicht gesetzt!")
    
    # Direkten Login-Endpoint definieren - make route path and function name match
    @app.route('/google_login')
    def google_login():
        """Eine direkte Google-Login-Route"""
        # Erstelle eine URL für die Google OAuth-Seite
        redirect_uri = url_for('google_callback', _external=True)
        print(f"Redirect URI being used: {redirect_uri}")  # Add this debug line
        
        google_auth_url = "https://accounts.google.com/o/oauth2/v2/auth"
        params = {
            'client_id': os.getenv('GOOGLE_CLIENT_ID'),
            'redirect_uri': redirect_uri,
            'response_type': 'code',
            'scope': 'email profile',
            'access_type': 'online',
            'state': 'direct_test'
        }
        
        auth_url = f"{google_auth_url}?{'&'.join([f'{k}={v}' for k, v in params.items()])}"
        return redirect(auth_url)

    # Callback für den direkten Login
    @app.route('/google_callback')
    def google_callback():
        """Callback für den direkten Google-Login"""
        code = request.args.get('code')
        
        if not code:
            flash('Login fehlgeschlagen (kein Code erhalten).', 'danger')
            return redirect(url_for('login'))
        
        # Code gegen Token tauschen
        try:
            token_url = "https://oauth2.googleapis.com/token"
            token_data = {
                'code': code,
                'client_id': os.getenv('GOOGLE_CLIENT_ID'),
                'client_secret': os.getenv('GOOGLE_CLIENT_SECRET'),
                'redirect_uri': url_for('google_callback', _external=True),
                'grant_type': 'authorization_code'
            }
            
            token_response = requests.post(token_url, data=token_data)
            
            if not token_response.ok:
                flash(f'Token-Abruf fehlgeschlagen: {token_response.text}', 'danger')
                return redirect(url_for('login'))
            
            token_info = token_response.json()
            access_token = token_info.get('access_token')
            
            # Benutzerinfo abrufen
            userinfo_url = "https://www.googleapis.com/oauth2/v2/userinfo"
            headers = {'Authorization': f'Bearer {access_token}'}
            userinfo_response = requests.get(userinfo_url, headers=headers)
            
            if not userinfo_response.ok:
                flash(f'Userinfo-Abruf fehlgeschlagen: {userinfo_response.text}', 'danger')
                return redirect(url_for('login'))
            
            user_info = userinfo_response.json()
            
            # Session aktualisieren
            email = user_info.get('email')
            name = user_info.get('name')
            
            # Setze die Session-Variablen
            user_id = session.get('user_id')
            if not user_id:
                user_id = str(uuid.uuid4())
                session['user_id'] = user_id
                
            session['email'] = email
            session['google_user_email'] = email
            session['user_name'] = name
            session['is_logged_via_google'] = True
            
            # Seller ID abrufen, wenn eine E-Mail vorhanden ist
            if email:
                seller_id = get_user_id_from_email(email)
                session['seller_id'] = seller_id
                if seller_id:
                    print(f"Seller ID gefunden und gesetzt: {seller_id}")
                else:
                    print(f"Keine Seller ID für E-Mail {email} gefunden")
            
            session.modified = True
            
            flash('Login erfolgreich!', 'success')
            return redirect(url_for('chat'))
        
        except Exception as e:
            print(f"Fehler beim Google-Login: {str(e)}")
            import traceback
            traceback.print_exc()
            flash(f'Fehler bei der Anmeldung: {str(e)}', 'danger')
            return redirect(url_for('login'))
    
    # Hilfsfunktion zur OAuth-Diagnose
    @app.route('/debug_oauth')
    def debug_oauth():
        """Diagnosewerkzeug für die OAuth-Konfiguration"""
        output = "<h1>OAuth Debug Info</h1>"
        
        # Umgebungsvariablen überprüfen (ohne Secrets zu zeigen)
        client_id = os.getenv('GOOGLE_CLIENT_ID', 'Nicht gesetzt')
        client_secret_status = "Gesetzt" if os.getenv('GOOGLE_CLIENT_SECRET') else "Nicht gesetzt"
        
        output += f"<p>GOOGLE_CLIENT_ID: {client_id[:5]}...{client_id[-5:] if len(client_id) > 10 else ''}</p>"
        output += f"<p>GOOGLE_CLIENT_SECRET: {client_secret_status}</p>"
        
        # Session-Status
        output += "<h2>Session Status</h2>"
        session_values = {k: v for k, v in session.items()}
        output += f"<pre>{json.dumps(session_values, indent=2)}</pre>"
        
        # Test Links - FIXED THIS LINE TO USE THE CORRECT FUNCTION NAME
        output += "<h2>Test Links</h2>"
        output += f"<p><a href='{url_for('google_login')}'>Google Login</a></p>"
        
        return output
    
    return app


app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your_default_secret_key')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# Sicherheitskonfiguration der Session
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Google OAuth konfigurieren
app = configure_google_auth(app)
Session(app)

logging.basicConfig(level=logging.info)

UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'uploaded_files')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# CSRF-Schutz
csrf = CSRFProtect(app)

openai.api_key = os.getenv('OPENAI_API_KEY')
if not openai.api_key:
    raise ValueError("Der OpenAI API-Schlüssel ist nicht gesetzt.")

###########################################
# Pinecone-Initialisierung (gRPC)
###########################################
pinecone_api_key = os.getenv('PINECONE_API_KEY')
pinecone_env = os.getenv('PINECONE_ENV')  # z.B. 'us-east-1'
pinecone_index_name = os.getenv('PINECONE_INDEX_NAME')

if not pinecone_api_key or not pinecone_index_name:
    raise ValueError("Bitte Pinecone API-Key und INDEX_NAME in .env setzen.")

# Pinecone-Client erstellen
pc = Pinecone(api_key=pinecone_api_key)
# Du kannst pinecone_env hier ggf. nicht explizit angeben – 
# Pinecone gRPC nutzt 'us-west1-gcp' oder 'us-east-1' ggf. 
# laut Projekt-Einstellungen.

# Index ggf. anlegen (Dimension=1536 für text-embedding-ada-002)
if not pc.has_index(pinecone_index_name):
    pc.create_index(
        name=pinecone_index_name,
        dimension=3072,
        metric="cosine",
        spec=ServerlessSpec(
            cloud="aws",
            region=pinecone_env or "us-east-1"
        )
    )
# Optional: Warten, bis Index ready
while not pc.describe_index(pinecone_index_name).status["ready"]:
    time.sleep(1)

# Index-Handle holen
index = pc.Index(pinecone_index_name)

###########################################
# Google Cloud Storage + Sonstige Einstellungen
###########################################
service_account_path = '/home/PfS/service_account_key.json'
if not os.path.exists(service_account_path):
    raise FileNotFoundError(f"Service Account Datei nicht gefunden: {service_account_path}")

credentials = service_account.Credentials.from_service_account_file(service_account_path)
client = storage.Client(credentials=credentials)
bucket_name = 'wissensbasis'
bucket = client.bucket(bucket_name)
wissensbasis_blob_name = 'wissensbasis.json'

DEBUG_CATEGORIES = {
    "API Calls": True,
    "Wissensbasis Download/Upload": True,
    "Bearbeiten von Einträgen": False,
    "Verschieben von Einträgen": False,
    "Sortieren von Einträgen": False,
    "UI Aktionen": False
}


themen_datei = '/home/PfS/themen.txt'

#############################
# Neuer Ordner für Chat-Logs
#############################
CHATLOG_FOLDER = 'chatlogs'
if not os.path.exists(CHATLOG_FOLDER):
    os.makedirs(CHATLOG_FOLDER)

#############################
# Neuer Ordner für Feedback
#############################
FEEDBACK_FOLDER = 'feedback'
if not os.path.exists(FEEDBACK_FOLDER):
    os.makedirs(FEEDBACK_FOLDER)

def store_chatlog(user_name, chat_history):
    """
    Speichert den Chatverlauf als Textdatei in CHATLOG_FOLDER.
    Dateiname enthält nur den user_name und das Datum (ohne Uhrzeit),
    damit alle Nachrichten vom selben Tag in derselben Datei gespeichert werden.
    """
    if not user_name:
        user_name = "Unbekannt"
    
    # Make sure the directory exists
    os.makedirs(CHATLOG_FOLDER, exist_ok=True)
    
    # Use just the date for the filename (not time) to group all chats from same day
    date_str = datetime.now().strftime("%Y-%m-%d")
    session_id = session.get('user_id', 'unknown')
    
    # Create a unique filename based on user, date and session ID
    filename = f"{user_name}_{date_str}_{session_id}.txt"
    filepath = os.path.join(CHATLOG_FOLDER, filename)
    
    # Write the complete chat history to the file (overwrite existing)
    with open(filepath, 'w', encoding='utf-8') as f:
        # Add a timestamp for this update
        current_time = datetime.now().strftime("%H:%M:%S")
        
        # Add a header
        f.write(f"=== CHAT LOG: {user_name} ===\n")
        f.write(f"Date: {date_str}\n")
        f.write(f"Session ID: {session_id}\n")
        f.write(f"Last Updated: {current_time}\n\n")
            
        # Write the entire chat history
        if chat_history:
            for idx, message in enumerate(chat_history, 1):
                user_msg = message.get('user', '').replace('\n', ' ')
                bot_msg = message.get('bot', '').replace('\n', ' ')
                
                f.write(f"Message {idx}:\n")
                f.write(f"  User: {user_msg}\n")
                f.write(f"  Bot : {bot_msg}\n\n")

def store_feedback(feedback_type, comment, chat_history, rated_message=""):
    # Make sure the feedback directory exists
    os.makedirs(FEEDBACK_FOLDER, exist_ok=True)
    
    name_in_session = session.get('user_name', 'Unbekannt')
    timestamp_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"{name_in_session}_{feedback_type}_{timestamp_str}.txt"
    filepath = os.path.join(FEEDBACK_FOLDER, filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"Feedback-Typ: {feedback_type}\n")
        f.write(f"Zeitpunkt: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
        f.write(f"Benutzer: {name_in_session}\n\n")
        
        # Add the specific message being rated
        if rated_message:
            f.write(f"Bewertete Nachricht:\n{rated_message}\n\n")
            
        f.write(f"Kommentar:\n{comment}\n\n")
        f.write("----- Chatverlauf -----\n\n")
        for idx, convo in enumerate(chat_history):
            user_msg = convo.get('user', '').replace('\n', ' ')
            bot_msg = convo.get('bot', '').replace('\n', ' ')
            f.write(f"Nachricht {idx+1}:\n")
            f.write(f"  User: {user_msg}\n")
            f.write(f"  Bot : {bot_msg}\n\n")

###########################################
# Chat-Statistik
###########################################
def calculate_chat_stats():
    total_count = 0
    year_count = 0
    month_count = 0
    day_count = 0

    now = datetime.now()
    current_year = now.year
    current_month = now.month
    current_day = now.day

    for filename in os.listdir(CHATLOG_FOLDER):
        if not filename.endswith(".txt"):
            continue
        
        filepath = os.path.join(CHATLOG_FOLDER, filename)
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Anzahl Chatnachrichten (Anzahl "User:" Vorkommen)
        message_pairs = content.count("User:")
        total_count += message_pairs


        try:
            date_str = filename.split("_")[1]  # => "YYYY-mm-dd"
            y, m, d = date_str.split("-")
            y, m, d = int(y), int(m), int(d)
        except:
            continue

        if y == current_year:
            year_count += message_pairs
            if m == current_month:
                month_count += message_pairs
                if d == current_day:
                    day_count += message_pairs

    return {
        'total': total_count,
        'year': year_count,
        'month': month_count,
        'today': day_count
    }

###########################################
# Download/Upload Wissensbasis (JSON)
###########################################
service_account_path = '/home/PfS/service_account_key.json'
if not os.path.exists(service_account_path):
    raise FileNotFoundError(f"Service Account Datei nicht gefunden: {service_account_path}")

credentials = service_account.Credentials.from_service_account_file(service_account_path)
client = storage.Client(credentials=credentials)
bucket_name = 'wissensbasis'
bucket = client.bucket(bucket_name)
wissensbasis_blob_name = 'wissensbasis.json'

def download_wissensbasis(max_retries=5, backoff_factor=1):
    debug_print("Wissensbasis Download/Upload", f"Versuche, Wissensbasis von '{wissensbasis_blob_name}' herunterzuladen.")
    blob = bucket.blob(wissensbasis_blob_name)

    content = '{}'
    for attempt in range(1, max_retries + 1):
        try:
            if blob.exists():
                content = blob.download_as_text(encoding='utf-8')
                debug_print("Wissensbasis Download/Upload", "Wissensbasis erfolgreich heruntergeladen.")
                break
            else:
                debug_print("Wissensbasis Download/Upload", "Wissensbasis-Datei existiert nicht.")
                break
        except Exception as e:
            debug_print("Wissensbasis Download/Upload", f"Fehler: {e}")
            if attempt < max_retries:
                wait_time = backoff_factor * (2 ** (attempt - 1))
                debug_print("Wissensbasis Download/Upload", f"Retry {attempt}: Warte {wait_time}s.")
                time.sleep(wait_time)
            else:
                flash(f"Fehler beim Herunterladen der Wissensbasis: {e}", 'danger')
                break

    wissensbasis = json.loads(content)
    for thema, unterthemen in wissensbasis.items():
        for unterthema, details in unterthemen.items():
            normalized_details = {key.lower(): value for key, value in details.items()}
            unterthemen[unterthema] = normalized_details
            normalized_details.setdefault('beschreibung', '')
            normalized_details.setdefault('inhalt', [])
    return wissensbasis

def upload_wissensbasis(wissensbasis, max_retries=5, backoff_factor=1):
    debug_print("Wissensbasis Download/Upload", "Versuche, Wissensbasis hochzuladen.")
    blob = bucket.blob(wissensbasis_blob_name)
    for attempt in range(1, max_retries + 1):
        try:
            blob.upload_from_string(
                json.dumps(wissensbasis, ensure_ascii=False, indent=4),
                content_type='application/json'
            )
            debug_print("Wissensbasis Download/Upload", "Wissensbasis hochgeladen.")
            time.sleep(0.5)
            return
        except Exception as e:
            debug_print("Wissensbasis Download/Upload", f"Fehler: {e}")
            if attempt < max_retries:
                wait_time = backoff_factor * (2 ** (attempt - 1))
                debug_print("Wissensbasis Download/Upload", f"Retry {attempt}: Warte {wait_time}s.")
                time.sleep(wait_time)
            else:
                flash(f"Fehler beim Hochladen der Wissensbasis: {e}", 'danger')

###########################################
# Themen (themen.txt) laden/aktualisieren
###########################################
themen_datei = '/home/PfS/themen.txt'

def lese_themenhierarchie(dateipfad):
    if not os.path.exists(dateipfad):
        return {}
    with open(dateipfad, 'r', encoding='utf-8') as f:
        zeilen = f.readlines()

    themen_dict = {}
    aktuelles_thema = None
    for zeile in zeilen:
        zeile = zeile.strip()
        if not zeile:
            continue
        match_thema = re.match(r'Thema\s*\d+:\s*.*', zeile)
        if match_thema:
            thema_key = zeile
            themen_dict[thema_key] = {}
            aktuelles_thema = thema_key
            continue
        match_unterpunkt = re.match(r'(\d+[a-z]*)\)?\s*([^/]+)\s*(//\s*(.*))?$', zeile)
        if match_unterpunkt:
            punkt_nummer = match_unterpunkt.group(1)
            punkt_titel = match_unterpunkt.group(2).strip()
            punkt_beschreibung = match_unterpunkt.group(4).strip() if match_unterpunkt.group(4) else ""
            if aktuelles_thema:
                themen_dict[aktuelles_thema][punkt_nummer] = {
                    "title": punkt_titel,
                    "beschreibung": punkt_beschreibung
                }
    return themen_dict

def lade_themen():
    return lese_themenhierarchie(themen_datei)

def get_next_thema_number(themen_dict):
    numbers = []
    for thema in themen_dict.keys():
        match = re.match(r'Thema\s+(\d+):\s+.*', thema)
        if match:
            numbers.append(int(match.group(1)))
    return max(numbers) + 1 if numbers else 1

def aktualisiere_themen(themen_dict):
    def sort_key(k):
        match = re.match(r'(\d+)([a-z]*)', k)
        if match:
            num = int(match.group(1))
            suf = match.group(2)
            return (num, suf)
        return (0, k)
    with open(themen_datei, 'w', encoding='utf-8') as file:
        for thema, unterpunkte in themen_dict.items():
            file.write(f"{thema}\n")
            sorted_up = sorted(unterpunkte.items(), key=lambda x: sort_key(x[0]))
            for punkt_nummer, punkt_info in sorted_up:
                punkt_titel = punkt_info['title']
                punkt_beschreibung = punkt_info.get('beschreibung', '')
                if punkt_beschreibung:
                    file.write(f"{punkt_nummer}) {punkt_titel}\t\t\t// {punkt_beschreibung}\n")
                else:
                    file.write(f"{punkt_nummer}) {punkt_titel}\n")
            file.write("\n")


###########################################
# Notfall-LOG-Funktion
###########################################
def log_notfall_event(user_id, notfall_art, user_message):
    log_path = "notfall_logs.json"
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "notfall_art": notfall_art,
        "message": user_message
    }
    logs = []
    if os.path.exists(log_path):
        with open(log_path, 'r', encoding='utf-8') as f:
            try:
                logs = json.load(f)
            except:
                logs = []
    logs.append(log_entry)
    with open(log_path, 'w', encoding='utf-8') as f:
        json.dump(logs, f, ensure_ascii=False, indent=2)

###########################################
# Wissenseintrag in JSON + Pinecone speichern
###########################################
def speichere_wissensbasis(eintrag):
    wissensbasis = download_wissensbasis()
    thema = eintrag.get("thema", "").strip()
    unterthema_full = eintrag.get("unterthema", "").strip()
    beschreibung = eintrag.get("beschreibung", "").strip()
    inhalt = eintrag.get("inhalt", "").strip()

    if thema and unterthema_full:
        match_unterthema = re.match(r'(\d+[a-z]*)\)?\s*(.*)', unterthema_full)
        if match_unterthema:
            unterthema_key = match_unterthema.group(1)
            unterthema_title = match_unterthema.group(2)
        else:
            unterthema_key = unterthema_full
            unterthema_title = ""

        unterthema_full_key = f"{unterthema_key}) {unterthema_title}"

        if thema not in wissensbasis:
            wissensbasis[thema] = {}
        if unterthema_full_key not in wissensbasis[thema]:
            wissensbasis[thema][unterthema_full_key] = {
                "beschreibung": beschreibung,
                "inhalt": []
            }
        if beschreibung:
            wissensbasis[thema][unterthema_full_key]["beschreibung"] = beschreibung
        if inhalt:
            wissensbasis[thema][unterthema_full_key]["inhalt"].append(inhalt)

        debug_print("Bearbeiten von Einträgen", f"Eintrag hinzugefügt/aktualisiert: {eintrag}")
        upload_wissensbasis(wissensbasis)

        # Pinecone: Upsert
        doc_id = f"{thema}-{unterthema_full_key}-{uuid.uuid4()}"
       
    else:
        flash("Thema und Unterthema müssen angegeben werden.", 'warning')



###########################################
# Kontakt mit OpenAI
###########################################
def contact_openai(messages, model=None):
    model = 'o3-mini'  # Changed from o3-mini-preview to o3-mini to match the model used in the chat route
    debug_print("API Calls", "contact_openai wurde aufgerufen – jetzt auf o3-mini gesetzt.")
    try:
        # Create the tool definitions first
        tools = create_function_definitions()
        
        # Add explicit tool_choice parameter to guide the model to use functions
        response = openai.chat.completions.create(
            model=model, 
            messages=messages,
            tools=tools,  # Add the tools parameter
            tool_choice="auto"  # Auto lets the model decide when to use functions
        )
        
        if response and response.choices:
            assistant_message = response.choices[0].message
            antwort_content = assistant_message.content.strip() if assistant_message.content else ""
            debug_print("API Calls", f"Antwort von OpenAI: {antwort_content}")

            # Check if the model chose to call a function
            tool_calls = assistant_message.tool_calls
            if tool_calls:
                debug_print("API Calls", f"Function Calls erkannt: {tool_calls}")
                for tool_call in tool_calls:
                    function_name = tool_call.function.name
                    function_args = json.loads(tool_call.function.arguments)
                    debug_print("API Calls", f"Funktion vom LLM gewählt: {function_name}, Argumente: {function_args}")

            return antwort_content, tool_calls  # Return both the content and tool calls
        else:
            antwort_content = "Keine Antwort erhalten."
            debug_print("API Calls", antwort_content)
            return antwort_content, None
        
    except Exception as e:
        debug_print("API Calls", f"Fehler: {e}")
        flash(f"Ein Fehler ist aufgetreten: {e}", 'danger')
        return None, None

def count_tokens(messages, model=None):
    model = 'o3-mini'
    try:
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        encoding = tiktoken.get_encoding("cl100k_base")
    token_count = 0
    for msg in messages:
        token_count += len(encoding.encode(msg['content']))
        token_count += 4
    token_count += 2
    return token_count


###########################################
# BigQuery Datenbezug
###########################################

# Outsourced to bigquery_functions.py

####################################
# EXTRACT INFOS FROM USER QUERY
####################################

# Outsourced to extract.py

####################################
# Function calling/ Tool Usage
####################################
import json
import logging
from openai import OpenAI

client = OpenAI()  # Annahme: dein API-Key ist in der Umgebung gesetzt

# Outsourced to tool_manager.py:
#def load_tool_descriptions()
#def create_tool_description_prompt()
#def select_tool()
#def select_optimal_tool_with_reasoning()
#def load_tool_config()



def format_simple_results(data_list):
    """Formatiert einfache Ergebnisse für Fallback-Antworten"""
    if not data_list:
        return "Keine Daten"
    
    result = []
    for item in data_list:
        try:
            # Versuche, häufig vorkommende Eigenschaften zu formatieren
            parts = []
            if "first_name" in item and "last_name" in item:
                parts.append(f"{item['first_name']} {item['last_name']}")
            if "agency_name" in item:
                parts.append(f"Agentur: {item['agency_name']}")
            if "bill_start" in item and "bill_end" in item:
                parts.append(f"Zeitraum: {format_date(item['bill_start'])} bis {format_date(item['bill_end'])}")
            if "prov_seller" in item:
                parts.append(f"Provision: {item['prov_seller']}€")
            
            if parts:
                result.append(" | ".join(parts))
            else:
                # Fallback: Zeige die ersten paar Eigenschaften
                simple_parts = []
                count = 0
                for key, value in item.items():
                    if count < 3 and key not in ["_id", "cs_id", "contract_id", "lead_id"]:
                        simple_parts.append(f"{key}: {value}")
                        count += 1
                result.append(" | ".join(simple_parts))
        except:
            # Bei Fehlern: Vereinfacht darstellen
            result.append(str(item)[:100] + "...")
    
    return "\n".join(result)

def create_function_definitions():
    """
    Erstellt die Function-Definitionen für OpenAI's Function-Calling basierend auf 
    den definierten Abfragemustern.
    
    Returns:
        list: Eine Liste von Function-Definitionen im Format, das von OpenAI erwartet wird
    """
    # Lade das JSON-Schema für Abfragemuster
    with open('query_patterns.json', 'r', encoding='utf-8') as f:
        query_patterns = json.load(f)
    
    tools = []
    
    # Erstelle für jedes Abfragemuster eine Funktion
    for query_name, query_info in query_patterns['common_queries'].items():
        function_def = {
            "type": "function",
            "function": {
                "name": query_name,
                "description": query_info['description'],
                "parameters": {
                    "type": "object",
                    "properties": {},
                    "required": query_info['required_parameters']
                }
            }
        }
        
        # Füge Parameter hinzu
        for param in query_info['required_parameters'] + query_info.get('optional_parameters', []):
            # Bestimme den Typ des Parameters basierend auf Namen (Heuristik)
            param_type = "string"
            if "id" in param:
                param_type = "string"
            elif "limit" in param or "count" in param or "_back" in param:
                param_type = "integer"
            elif "date" in param or "time" in param:
                param_type = "string"  # Datum als String, wird später konvertiert
            elif param == "contacted":
                param_type = "boolean"
            
            # Bestimme die Beschreibung des Parameters
            param_desc = f"Parameter {param} für die Abfrage"
            if param == "seller_id":
                param_desc = "Die ID des Verkäufers, dessen Daten abgefragt werden sollen"
            elif param == "lead_id":
                param_desc = "Die ID des Leads, dessen Daten abgefragt werden sollen"
            elif param == "limit":
                param_desc = "Maximale Anzahl der zurückzugebenden Datensätze"
            
            # Füge Parameter zur Funktionsdefinition hinzu
            function_def["function"]["parameters"]["properties"][param] = {
                "type": param_type,
                "description": param_desc
            }
            
            # Füge Enumerationen für bestimmte Parameter hinzu
            if param == "ticketable_type":
                function_def["function"]["parameters"]["properties"][param]["enum"] = [
                    "Lead", "Contract", "CareStay", "Visor", "Posting"
                ]
        
        tools.append(function_def)
    
    return tools

def create_system_prompt(table_schema):
    # Bestehendes System-Prompt generieren
    prompt = "Du bist ein hilfreicher KI-Assistent, der bei der Verwaltung von Pflegedaten hilft."
    prompt += "\n\nDu hast Zugriff auf eine Datenbank mit folgenden Tabellen:\n"
    
    for table_name, table_info in table_schema.get("tables", {}).items():
        prompt += f"\n- {table_name}: {table_info.get('description', 'Keine Beschreibung')}"
        prompt += "\n  Felder:"
        for field_name, field_info in table_info.get("fields", {}).items():
            prompt += f"\n    - {field_name}: {field_info.get('description', 'Keine Beschreibung')}"
    
    # Ergänze das Prompt mit wichtigen Anweisungen zur Funktionsnutzung
    prompt += """
    
    KRITISCH WICHTIG: Du bist ein Assistent, der NIEMALS Fragen zu Datenbank-Daten direkt beantwortet!
    
    1. Bei JEDER Frage zu Care Stays, Verträgen, Leads oder anderen Daten MUSST du eine der bereitgestellten Funktionen verwenden.
    2. Ohne Funktionsaufruf hast du KEINEN Zugriff auf aktuelle Daten.
    3. Generiere NIEMALS Antworten aus eigenem Wissen, wenn die Information in der Datenbank zu finden ist.
    4. Bei zeitbezogenen Anfragen (z.B. "im Mai") nutze IMMER die Funktion get_care_stays_by_date_range.
    
    Dein Standardverhalten bei Datenabfragen:
    1. Analysiere die Nutzerfrage
    2. Wähle die passende Funktion
    3. Rufe die Funktion mit korrekten Parametern auf
    4. Warte auf das Ergebnis
    5. Nutze dieses Ergebnis für deine Antwort
    """
    
    return prompt

def stream_text_response(response_text, user_message, session_data):
    """
    Generiert einen Stream für direkte Textantworten (z.B. Wissensbasis-Antworten)
    """
    try:
        # Debug-Events für die Verbindungsdiagnose
        yield f"data: {json.dumps({'type': 'debug', 'message': 'Stream-Start (Text Response)'})}\n\n"
        
        # Stream-Start
        yield f"data: {json.dumps({'type': 'start'})}\n\n"
        
        # Teile die Antwort in Chunks auf für ein natürlicheres Streaming-Gefühl
        chunk_size = 15  # Anzahl der Wörter pro Chunk
        words = response_text.split()
        
        for i in range(0, len(words), chunk_size):
            chunk = " ".join(words[i:i+chunk_size])
            yield f"data: {json.dumps({'type': 'text', 'content': chunk + ' '})}\n\n"
            time.sleep(0.05)  # Kleine Verzögerung für natürlichere Ausgabe
        
        # Stream beenden
        yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': response_text})}\n\n"
        yield f"data: {json.dumps({'type': 'debug', 'message': 'Stream complete with function execution'})}\n\n"
        yield f"data: {json.dumps({'type': 'end'})}\n\n"
    except Exception as e:
        logging.exception("Fehler im Text-Stream")
        yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler: {str(e)}'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': 'Es ist ein Fehler aufgetreten.'})}\n\n"
        yield f"data: {json.dumps({'type': 'end'})}\n\n"

def generate_conversational_clarification_stream(clarification_data):
    """
    Generiert einen Stream für konversationelle Rückfragen ohne Buttons
    """
    try:
        yield f"data: {json.dumps({'type': 'clarification_start'})}\n\n"
        
        # Sende die Rückfrage als normalen Text
        message = clarification_data.get("clarification_message", "Bitte präzisiere deine Anfrage")
        yield f"data: {json.dumps({'type': 'text', 'content': message})}\n\n"
        
        # Wir verwenden keine Button-Optionen mehr, sondern erwarten eine natürliche Antwort
        yield f"data: {json.dumps({'type': 'complete', 'user': clarification_data.get('original_question', ''), 'bot': message})}\n\n"
        yield f"data: {json.dumps({'type': 'conversational_clarification_mode'})}\n\n"
        
        yield f"data: {json.dumps({'type': 'clarification_end'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler bei Rückfrage: {str(e)}'})}\n\n"

def generate_clarification_stream(human_in_loop_data):
    """
    Generiert einen Stream für text-basierte Rückfragen (ohne Buttons)
    """
    try:
        yield f"data: {json.dumps({'type': 'text_clarification_start'})}\n\n"
        
        # Send the clarification question as regular text - no need for buttons
        message = human_in_loop_data.get("message", "Bitte präzisiere deine Anfrage")
        
        # Split the options out of the message if present
        # This ensures the client can directly display the message text
        yield f"data: {json.dumps({'type': 'text', 'content': message})}\n\n"
        
        # Signal that this is the end of this message - client will render it as normal text
        yield f"data: {json.dumps({'type': 'text_clarification_end'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler bei Rückfrage: {str(e)}'})}\n\n"

def stream_response(messages, tools, tool_choice, seller_id, extracted_args, user_message, session_data):
    """Stream the OpenAI response and handle function calls within the stream
    
    Args:
        messages: OpenAI message array
        tools: Available tools/functions
        tool_choice: Which tool to use
        seller_id: The seller's ID
        extracted_args: Any extracted date parameters
        user_message: The original user message
        session_data: A dictionary containing all needed session data
    """
    # Extract session data - this avoids accessing session in the generator
    user_id = session_data["user_id"]
    user_name = session_data["user_name"]
    chat_key = session_data["chat_key"]
    chat_history = session_data["chat_history"]
    
    try:
        # Debug-Events für die Verbindungsdiagnose
        yield f"data: {json.dumps({'type': 'debug', 'message': 'Stream-Start'})}\n\n"
        #yield f"data: {json.dumps({'type': 'text', 'content': 'Test-Content vom Server'})}\n\n"

        debug_print("API Calls", f"Streaming-Anfrage an OpenAI mit Function Calling")
        response = openai.chat.completions.create(
            model="o3-mini",
            messages=messages,
            tools=tools,
            tool_choice=tool_choice,
            stream=True  # Enable streaming
        )
        
        initial_response = ""
        function_calls_data = []
        has_function_calls = False
        
        # Stream the initial response
        yield f"data: {json.dumps({'type': 'start'})}\n\n"
        
        for chunk in response:
            if chunk.choices[0].delta.content:
                text_chunk = chunk.choices[0].delta.content
                initial_response += text_chunk
                yield f"data: {json.dumps({'type': 'text', 'content': text_chunk})}\n\n"
            
            # Check if this chunk contains function call info
            if hasattr(chunk.choices[0].delta, 'tool_calls') and chunk.choices[0].delta.tool_calls:
                has_function_calls = True
                # Collect function call data
                for tool_call in chunk.choices[0].delta.tool_calls:
                    tool_index = tool_call.index
                    
                    # Initialize the tool call if needed
                    while len(function_calls_data) <= tool_index:
                        function_calls_data.append({"id": str(tool_index), "name": "", "args": ""})
                    
                    if hasattr(tool_call, 'function'):
                        if hasattr(tool_call.function, 'name') and tool_call.function.name:
                            function_calls_data[tool_index]["name"] = tool_call.function.name
                        
                        if hasattr(tool_call.function, 'arguments') and tool_call.function.arguments:
                            function_calls_data[tool_index]["args"] += tool_call.function.arguments
        
        # If function calls detected, execute them
        if has_function_calls:
            yield f"data: {json.dumps({'type': 'function_call_start'})}\n\n"
            yield f"data: {json.dumps({'type': 'debug', 'message': f'Detected {len(function_calls_data)} function calls'})}\n\n"
            
            function_responses = []
            for func_data in function_calls_data:
                if func_data["name"] and func_data["args"]:
                    try:
                        function_name = func_data["name"]
                        function_args = json.loads(func_data["args"])
                        
                        # Add seller_id and extracted date parameters
                        if seller_id:
                            function_args["seller_id"] = seller_id
                        
                        for key, value in extracted_args.items():
                            if key not in function_args or not function_args[key]:
                                function_args[key] = value
                        
                        # Execute the function
                        debug_print("Function", f"Streaming: Executing {function_name} with args {function_args}")
                        yield f"data: {json.dumps({'type': 'debug', 'message': f'Executing function {function_name}'})}\n\n"
                        
                        function_response = handle_function_call(function_name, function_args)
                        
                        # Add to function responses
                        function_responses.append({
                            "role": "tool",
                            "tool_call_id": func_data["id"],
                            "content": function_response
                        })
                        
                        yield f"data: {json.dumps({'type': 'function_result', 'name': function_name})}\n\n"
                        yield f"data: {json.dumps({'type': 'debug', 'message': f'Function executed successfully'})}\n\n"
                        
                    except Exception as e:
                        debug_print("Function", f"Error executing function: {str(e)}")
                        yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler bei Funktionsausführung: {str(e)}'})}\n\n"
            
            # Second call to get final response
            if function_responses:
                # Properly format the tool_calls with the required 'type' field
                formatted_tool_calls = []
                for func_data in function_calls_data:
                    formatted_tool_calls.append({
                        "type": "function",  # This is the required field
                        "id": func_data["id"],
                        "function": {
                            "name": func_data["name"],
                            "arguments": func_data["args"]
                        }
                    })
                
                # Create the second messages with properly formatted tool_calls
                second_messages = messages + [{
                    "role": "assistant", 
                    "content": initial_response,
                    "tool_calls": formatted_tool_calls
                }] + function_responses

                # Debug vor dem zweiten API-Call
                yield f"data: {json.dumps({'type': 'debug', 'message': 'Preparing for second API call'})}\n\n"
                
                # WICHTIG: Final response start event - stelle sicher, dass es gesendet wird
                yield f"data: {json.dumps({'type': 'final_response_start'})}\n\n"
                
                # Initiiere den zweiten API-Call
                yield f"data: {json.dumps({'type': 'debug', 'message': 'Starting second API call'})}\n\n"
                
                try:
                    final_response = openai.chat.completions.create(
                        model="o3-mini",
                        messages=second_messages,
                        stream=True
                    )
                    
                    yield f"data: {json.dumps({'type': 'debug', 'message': 'Second API call initiated'})}\n\n"
                    
                    final_text = ""
                    for chunk in final_response:
                        if chunk.choices[0].delta.content:
                            text_chunk = chunk.choices[0].delta.content
                            final_text += text_chunk
                            yield f"data: {json.dumps({'type': 'text', 'content': text_chunk})}\n\n"
                    
                    # Vollständige Antwort zusammenstellen
                    full_response = initial_response + "\n\n" + final_text
                    
                    # Session aktualisieren und Stream beenden
                    yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': full_response})}\n\n"
                    yield f"data: {json.dumps({'type': 'debug', 'message': 'Stream complete with function execution'})}\n\n"
                    yield f"data: {json.dumps({'type': 'end'})}\n\n"
                    
                except Exception as e:
                    debug_print("API Calls", f"Error in second call: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler beim zweiten API-Call: {str(e)}'})}\n\n"
                    # Trotz Fehler die Session aktualisieren
                    yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': initial_response})}\n\n"
                    yield f"data: {json.dumps({'type': 'end'})}\n\n"
            else:
                # No valid function responses, just save initial response
                yield f"data: {json.dumps({'type': 'debug', 'message': 'Function execution produced no valid responses'})}\n\n"
                yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': initial_response})}\n\n"
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
        else:
            # No function calls, just save initial response
            yield f"data: {json.dumps({'type': 'debug', 'message': 'No function calls detected'})}\n\n"
            yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': initial_response})}\n\n"
            yield f"data: {json.dumps({'type': 'end'})}\n\n"
    
    except Exception as e:
        logging.exception("Fehler im Stream")
        yield f"data: {json.dumps({'type': 'error', 'content': f'Fehler: {str(e)}'})}\n\n"
        # Auch bei Fehler versuchen, Stream ordnungsgemäß zu beenden
        yield f"data: {json.dumps({'type': 'complete', 'user': user_message, 'bot': 'Es ist ein Fehler aufgetreten.'})}\n\n"
        yield f"data: {json.dumps({'type': 'end'})}\n\n"    




@app.route("/clarify", methods=["POST"])
def handle_clarification():
    """
    Verarbeitet die Antwort auf eine Rückfrage - unterstützt sowohl Legacy Button-Modus als auch
    neue konversationelle Rückfragen
    """
    try:
        # Stellen wir sicher, dass eine Benutzer-Session existiert
        if not session.get("user_id"):
            logging.warning("Clarification-Anfrage ohne aktive Benutzer-Session")
            return jsonify({"error": "Keine aktive Benutzer-Session gefunden"}), 400
        
        # Für AJAX-Anfragen
        is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"
        
        # Debug-Informationen
        logging.info(f"Verarbeite Clarification-Antwort, AJAX: {is_ajax}")
        
        # Prüfen, ob wir im konversationellen oder Button-Modus sind
        conversational_mode = session.get("clarification_in_progress", False)
        
        if conversational_mode:
            # Konversationeller Modus: Antwort als natürlichen Text behandeln
            user_message = request.form.get("message", "").strip()
            
            if not user_message:
                error_msg = "Keine Antwort auf die Rückfrage angegeben"
                logging.error(error_msg)
                if is_ajax:
                    return jsonify({"error": error_msg}), 400
                flash(error_msg, "warning")
                return redirect(url_for("chat"))
            
            # Benutzerantwort wird in der normalen Chat-Route verarbeitet
            # Wir senden eine Erfolgsmeldung zurück, da die Verarbeitung im Chat-Handler erfolgt
            if is_ajax:
                return jsonify({"success": True, "message": "Antwort wird verarbeitet"})
            return redirect(url_for("chat"))
        
        # Legacy Button-Modus
        try:
            # Lese die Option aus dem Formular
            selected_option_index = int(request.form.get("option_index", "0"))
            logging.info(f"Ausgewählter Options-Index: {selected_option_index}")
        except ValueError:
            error_msg = "Ungültiger Options-Index Format (keine Zahl)"
            logging.error(f"Fehler beim Parsen des Options-Index: {request.form.get('option_index')}")
            if is_ajax:
                return jsonify({"error": error_msg}), 400
            flash(error_msg, "danger")
            return redirect(url_for("chat"))
            
        # Prüfe, ob die Human-in-Loop-Daten in der Session vorhanden sind
        human_in_loop_data = session.get("human_in_loop_data")
        
        # Logge die vorhandenen Session-Daten für Debugging
        session_keys = list(session.keys())
        logging.debug(f"Verfügbare Session-Keys: {session_keys}")
        
        if not human_in_loop_data:
            error_msg = "Keine Rückfrage-Daten in der Session gefunden"
            logging.error(error_msg)
            if is_ajax:
                return jsonify({"error": error_msg}), 400
            flash(error_msg, "danger")
            return redirect(url_for("chat"))
            
        if "options" not in human_in_loop_data:
            error_msg = "Ungültige Rückfrage-Daten: Keine Optionen vorhanden"
            logging.error(f"human_in_loop_data ohne 'options': {human_in_loop_data}")
            if is_ajax:
                return jsonify({"error": error_msg}), 400
            flash(error_msg, "danger")
            return redirect(url_for("chat"))
            
        options = human_in_loop_data.get("options", [])
        logging.info(f"Anzahl verfügbarer Optionen: {len(options)}")
        
        # Prüfe, ob der Index gültig ist
        if selected_option_index < 0 or selected_option_index >= len(options):
            error_msg = f"Ungültiger Options-Index: {selected_option_index} (max: {len(options)-1})"
            logging.error(error_msg)
            if is_ajax:
                return jsonify({"error": error_msg}), 400
            flash(error_msg, "danger")
            return redirect(url_for("chat"))
            
        # Hole ausgewählte Option
        selected_option = options[selected_option_index]
        logging.info(f"Ausgewählte Option: {selected_option.get('text')} -> {selected_option.get('query')}")
        
        # Verarbeiten von Parametern, speziell für den Kunden Küll
        params = selected_option.get("params", {})
        if params and "customer_name" in params:
            customer_name = params["customer_name"]
            if customer_name.lower() in ["küll", "kull", "kühl", "kuehl", "kuell"]:
                logging.info("Normalisiere Kunden-Parameter 'Küll'")
                params["customer_name"] = "Küll"
                selected_option["params"] = params

        # Extrahiere die benötigten Informationen
        selected_query = selected_option.get("query")
        selected_params = selected_option.get("params", {})
        
        # Verarbeitung direkt hier für Kundenabfragen
        if selected_query == "get_customer_history" and "customer_name" in selected_params:
            try:
                # Kundenname extrahieren
                customer_name = selected_params.get("customer_name")
                
                # seller_id aus der Session hinzufügen
                if "seller_id" in session:
                    selected_params["seller_id"] = session.get("seller_id")
                
                logging.info(f"Führe direkt get_customer_history für '{customer_name}' aus")
                
                # Funktion direkt aufrufen
                result = handle_function_call(selected_query, selected_params)
                result_data = json.loads(result)
                
                # Ergebnis formatieren und den Chat-Verlauf aktualisieren
                if "data" in result_data and len(result_data["data"]) > 0:
                    formatted_result = format_customer_details(result_data)
                    
                    # Original-Anfrage holen
                    original_request = session.get("human_in_loop_original_request", "")
                    
                    # Chat-History aktualisieren
                    chat_key = f"chat_history_{session.get('user_id')}"
                    chat_history = session.get(chat_key, [])
                    
                    # Benutzeranfrage und Antwort hinzufügen
                    chat_history.append({"role": "user", "content": original_request})
                    chat_history.append({"role": "assistant", "content": formatted_result})
                    session[chat_key] = chat_history
                    
                    # Erfolg zurückgeben
                    if is_ajax:
                        return jsonify({
                            "success": True, 
                            "response": formatted_result,
                            "message": "Kundeninformationen erfolgreich abgerufen."
                        })
                    else:
                        # Bei normaler Anfrage: Flash-Nachricht und Weiterleitung
                        flash("Kundeninformationen erfolgreich abgerufen.", "success")
                        # WICHTIG: Antwort in der Session speichern für die Anzeige
                        session["last_response"] = formatted_result
                        session.modified = True
                        return redirect(url_for("chat"))
                else:
                    # Keine Daten gefunden
                    error_msg = f"Keine Daten für Kunde '{customer_name}' gefunden."
                    if is_ajax:
                        return jsonify({"error": error_msg})
                    flash(error_msg, "warning")
                    return redirect(url_for("chat"))
            
            except Exception as e:
                logging.error(f"Fehler bei direkter Kundenanfrage: {str(e)}")
                if is_ajax:
                    return jsonify({"error": str(e)})
                flash(f"Fehler bei der Verarbeitung: {str(e)}", "danger")
                return redirect(url_for("chat"))
        
        # Für andere Anfragen: in Session speichern für nächsten Request
        session["human_in_loop_clarification_response"] = selected_option
        
        # Speichere den original request für die Kontext-Kontinuität
        if "human_in_loop_original_request" in session:
            original_request = session.get("human_in_loop_original_request")
            session["pending_query"] = original_request
            logging.info(f"Original-Anfrage gespeichert: {original_request}")
        else:
            logging.warning("Keine Original-Anfrage in der Session gefunden")
        
        # Entferne die Human-in-the-Loop-Daten aus der Session
        session.pop("human_in_loop_data", None)
        session.pop("human_in_loop_original_request", None)
        
        # Stelle sicher, dass der Chatverlauf erhalten bleibt
        session.modified = True
        
        # Bei AJAX-Anfragen mehr Informationen zurückgeben für clientseitige Verarbeitung
        if is_ajax:
            logging.info("Sende AJAX-Erfolgsantwort mit Option")
            # Hier nehmen wir an, dass im nächsten Request eine Antwort erstellt wird
            # Daher senden wir ein Signal an die Client-Seite, dass es eine Anfrage gab,
            # die beim nächsten Request beantwortet wird (im GET handler der chat route)
            return jsonify({
                "success": True, 
                "message": "Option ausgewählt",
                "selected_option": selected_option.get("text", ""),
                "query": selected_query,
                "need_followup": True
            })
        
        # Ansonsten wie gehabt weiterleiten
        logging.info("Weiterleitung zur Chat-Seite")
        return redirect(url_for("chat"))
    except Exception as e:
        # Allgemeine Fehlerbehandlung
        error_msg = f"Fehler bei der Verarbeitung der Rückfrage: {str(e)}"
        logging.exception("Unerwarteter Fehler bei der Verarbeitung der Rückfrage")
        
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({"error": error_msg, "success": False}), 500
            
        flash(error_msg, "danger")
        return redirect(url_for("chat"))
        
    except Exception as e:
        logging.error(f"Fehler bei der Verarbeitung der Rückfrage: {e}")
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({"error": str(e)}), 500
        flash(f"Ein Fehler ist aufgetreten: {str(e)}", "danger")
        return redirect(url_for("chat"))

@app.route("/", methods=["GET", "POST"])
def chat():
    try:
        user_id = session.get("user_id")
        user_name = session.get("user_name")

        if not user_name:
            return redirect(url_for("google_login"))

        seller_id = session.get("seller_id")
        if not seller_id and session.get("email"):
            seller_id = get_user_id_from_email(session.get("email"))
            if seller_id:
                session["seller_id"] = seller_id
                session.modified = True

        chat_key = f"chat_history_{user_id}"
        if chat_key not in session:
            session[chat_key] = []
        chat_history = session[chat_key]

        # Einfache Anzeige der Chat-History
        display_chat_history = chat_history

        try:
            with open("table_schema.json", "r", encoding="utf-8") as f:
                table_schema = json.load(f)
            with open("query_patterns.json", "r", encoding="utf-8") as f:
                query_patterns = json.load(f)
        except Exception as e:
            logging.error(f"Fehler beim Laden der Schema-Dateien: {e}")
            table_schema = {"tables": {}}
            query_patterns = {"common_queries": {}}

        if request.method == "POST":
            user_message = request.form.get("message", "").strip()
            notfall_aktiv = request.form.get("notfallmodus") == "1"
            notfall_art = request.form.get("notfallart", "").strip()
            
            # Check if this is a streaming request
            stream_mode = request.form.get("stream", "0") == "1"

            if not user_message:
                flash("Bitte geben Sie eine Nachricht ein.", "warning")
                return (
                    jsonify({"error": "Bitte geben Sie eine Nachricht ein."}),
                    400,
                ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else redirect(
                    url_for("chat")
                )

            wissensbasis = download_wissensbasis()
            if not wissensbasis:
                flash("Die Wissensbasis konnte nicht geladen werden.", "danger")
                return (
                    jsonify({"error": "Die Wissensbasis konnte nicht geladen werden."}),
                    500,
                ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else redirect(
                    url_for("chat")
                )

            if notfall_aktiv:
                session["notfall_mode"] = True
                user_message = (
                    f"ACHTUNG NOTFALL - Thema 9: Notfälle & Vertragsgefährdungen.\n"
                    f"Ausgewählte Notfalloption(en): {notfall_art}\n\n"
                    + user_message
                )
                log_notfall_event(user_id, notfall_art, user_message)
            else:
                session.pop("notfall_mode", None)

            # Verbesserter System Prompt
            system_prompt = create_system_prompt(table_schema)
            prompt_wissensbasis_abschnitte = [
                f"Thema: {thema}, Unterthema: {unterthema_full}, Beschreibung: {details.get('beschreibung', '')}, Inhalt: {'. '.join(details.get('inhalt', []))}"
                for thema, unterthemen in wissensbasis.items()
                for unterthema_full, details in unterthemen.items()
            ]
            system_prompt += f"\n\nWissensbasis:\n{chr(10).join(prompt_wissensbasis_abschnitte)}"
            system_prompt = (
                f"Der Name deines Gesprächspartners lautet {user_name}.\n"
                + system_prompt
                + (f"\n\nDu sprichst mit einem Vertriebspartner mit der ID {seller_id}." if seller_id else "")
            )

            # Setup tools and message collections for different approaches
            tools = create_function_definitions()
            debug_print("API Setup", f"System prompt (gekürzt): {system_prompt[:200]}...")
            debug_print("API Setup", f"Anzahl definierter Tools: {len(tools)}")

            messages = [
                {"role": "developer", "content": system_prompt},
                {"role": "user", "content": user_message},
            ]

            # Sammle alle relevanten Session-Daten für den mehrstufigen Prozess
            session_data = {
                "user_id": user_id,
                "user_name": user_name,
                "seller_id": seller_id,
                "email": session.get("email"),
                "chat_key": chat_key,
                "chat_history": list(chat_history) if stream_mode else None  # Nur für Streaming benötigt
            }

            # Debug-Modus: Direkte Erzwingung eines bestimmten Tools
            debug_force = request.args.get("force_function")
            if debug_force:
                debug_print("DEBUG", f"Erzwinge Funktionsaufruf: {debug_force}")
                # Hier verwenden wir die alte Methode der direkten OpenAI-Aufrufe mit erzwungenem Tool
                tool_choice = {"type": "function", "function": {"name": debug_force}}
                use_legacy_approach = True
            else:
                # Für den normalen Betrieb verwenden wir den neuen mehrstufigen Ansatz
                use_legacy_approach = False

            try:
                # Streaming-Modus
                if stream_mode and request.headers.get("Accept") == "text/event-stream":
                    if use_legacy_approach:
                        # Legacy Streaming mit direkter Tool-Auswahl für Debug
                        return Response(
                            stream_response(messages, tools, tool_choice, seller_id, extract_date_params(user_message), user_message, session_data),
                            content_type="text/event-stream"
                        )
                    else:
                        # Vereinfachter Ansatz für die erste Implementierung
                        # Verwende den bestehenden select_optimal_tool_with_reasoning für die erste Version
                        tool_config = load_tool_config()
                        selected_tool, reasoning = select_optimal_tool_with_reasoning(user_message, tools, tool_config)
                        
                        # "Erfassungsbogen" oder ähnliche Anfragen direkt an die Wissensbasis weiterleiten
                        wissensbasis_keywords = ["wissensdatenbank", "wissensbasis", "erfassungsbogen", "handbuch", 
                                               "anleitung", "wie funktioniert", "erklär mir", "was ist", "was sind"]
                        
                        is_wissensbasis_query = False
                        user_message_lower = user_message.lower()
                        
                        for keyword in wissensbasis_keywords:
                            if keyword in user_message_lower:
                                is_wissensbasis_query = True
                                break
                        
                        if is_wissensbasis_query:
                            # Lade Wissensbasis und beantworte direkt
                            wissensbasis_data = download_wissensbasis()
                            
                            system_prompt = """
                            Du bist ein hilfreicher Assistent für ein Pflegevermittlungsunternehmen. 
                            Beantworte die Frage basierend auf der bereitgestellten Wissensbasis.
                            Sei klar, präzise und sachlich. Wenn du die Antwort nicht in der Wissensbasis findest, 
                            sage ehrlich, dass du es nicht weißt.
                            """
                            
                            wissensbasis_messages = [
                                {"role": "developer", "content": system_prompt},
                                {"role": "user", "content": f"Wissensbasis: {wissensbasis_data}\n\nFrage: {user_message}"}
                            ]
                            
                            response = openai.chat.completions.create(
                                model="o3-mini",  # Capabilities for knowledge-based queries
                                messages=wissensbasis_messages
                            )
                            
                            wissensbasis_response = response.choices[0].message.content
                            return Response(
                                stream_text_response(wissensbasis_response, user_message, session_data),
                                content_type="text/event-stream"
                            )
                        
                        # Prüfen, ob eine Rückfrage ausgelöst wurde
                        elif selected_tool == "human_in_loop_clarification":
                            # Neue text-basierte Rückfrage ohne Buttons
                            human_in_loop_data = session.get("human_in_loop_data")
                            
                            # Process the data differently based on type
                            if human_in_loop_data and human_in_loop_data.get("type") == "text_clarification":
                                # Process as a text clarification - just stream the message
                                debug_print("Clarification", "Streaming text-based clarification")
                                return Response(
                                    generate_clarification_stream(human_in_loop_data),
                                    content_type="text/event-stream"
                                )
                            else:
                                # Process as a legacy clarification for backward compatibility
                                debug_print("Clarification", "Streaming legacy clarification")
                                return Response(
                                    generate_clarification_stream(human_in_loop_data),
                                    content_type="text/event-stream"
                                )
                        
                        # Direkte Konversationsanfragen wie mathematische Berechnungen
                        elif selected_tool == "direct_conversation":
                            debug_print("Tool-Auswahl", "Direkter Konversationsmodus erkannt - Leite Anfrage direkt an LLM weiter")
                            direct_messages = [
                                {"role": "system", "content": "Du bist ein hilfreicher Assistent für ein Pflegevermittlungsunternehmen."},
                                {"role": "user", "content": user_message}
                            ]
                            
                            response = openai.chat.completions.create(
                                model="o3-mini",
                                messages=direct_messages
                            )
                            
                            direct_response = response.choices[0].message.content
                            return Response(
                                stream_text_response(direct_response, user_message, session_data),
                                content_type="text/event-stream"
                            )
                        
                        
                        
                        # Korrektes Format für tool_choice erstellen
                        tool_choice = {"type": "function", "function": {"name": selected_tool}} if selected_tool else "auto"
                                                
                        return Response(
                            stream_response(
                                messages, 
                                tools, 
                                tool_choice,  # Statt dem String das korrekt formatierte Objekt übergeben
                                seller_id, 
                                extract_enhanced_date_params(user_message), 
                                user_message, 
                                session_data
                            ),
                            content_type="text/event-stream"
                        )
                           
                        
                
                # Nicht-Streaming-Modus
                if use_legacy_approach:
                    # Legacy-Ansatz (für Debug-Modus)
                    debug_print("API Calls", f"Legacy-Ansatz mit explizitem Function Calling: {debug_force}")
                    response = openai.chat.completions.create(
                        model="o3-mini",
                        messages=messages,
                        tools=tools,
                        tool_choice=tool_choice
                    )
                    
                    assistant_message = response.choices[0].message
                    debug_print("API Response", f"Message type: {type(assistant_message)}")
                    debug_print("API Response", f"Has tool calls: {hasattr(assistant_message, 'tool_calls') and bool(assistant_message.tool_calls)}")

                    # Initialize the antwort variable
                    antwort = None

                    if assistant_message.tool_calls:
                        debug_print("API Calls", f"Function Calls erkannt: {assistant_message.tool_calls}")
                        function_responses = []

                        for tool_call in assistant_message.tool_calls:
                            function_name = tool_call.function.name
                            function_args = json.loads(tool_call.function.arguments)
                            debug_print("Function", f"Name: {function_name}, Argumente vor Modifikation: {function_args}")

                            # Standardargumente hinzufügen und überschreiben
                            if seller_id:
                                function_args["seller_id"] = seller_id
                            
                            # Nur extrahierte Datumswerte hinzufügen, wenn sie nicht bereits gesetzt sind
                            extracted_args = extract_enhanced_date_params(user_message)
                            for key, value in extracted_args.items():
                                if key not in function_args or not function_args[key]:
                                    function_args[key] = value

                            debug_print("Function", f"Argumente nach Modifikation: {function_args}")
                            function_response = handle_function_call(function_name, function_args)
                            
                            # Parsen der Function Response für bessere Logging
                            try:
                                parsed_response = json.loads(function_response)
                                response_status = parsed_response.get("status", "unknown")
                                response_count = len(parsed_response.get("data", [])) if "data" in parsed_response else 0
                                debug_print("Function Response", f"Status: {response_status}, Anzahl Ergebnisse: {response_count}")
                            except:
                                debug_print("Function Response", "Konnte Response nicht parsen")
                            
                            function_responses.append(
                                {
                                    "role": "tool",
                                    "tool_call_id": tool_call.id,
                                    "content": function_response,
                                }
                            )

                        # Second call to OpenAI with function results
                        second_messages = messages + [assistant_message.model_dump(exclude_unset=True)] + function_responses
                        debug_print("API Calls", f"Zweiter Aufruf an OpenAI mit {len(function_responses)} Funktionsantworten")
                        second_response = openai.chat.completions.create(model="o3-mini", messages=second_messages)
                        final_message = second_response.choices[0].message
                        antwort = final_message.content
                        debug_print("API Calls", f"Finale Antwort: {antwort[:100]}...")
                        # Warnung ins Log schreiben, wenn keine Funktion aufgerufen wurde
                        if any(term in user_message.lower() for term in ["care", "pflege", "kunden", "verträge", "mai", "monat"]):
                            logging.warning(f"Keine Funktion aufgerufen trotz relevanter Anfrage: '{user_message}'")
                
                else:
                    # Neuer mehrstufiger Ansatz
                    debug_print("API Calls", f"Verwende mehrstufigen Ansatz für: {user_message}")
                    antwort = process_user_query(user_message, session_data)
                    debug_print("API Calls", f"Mehrstufige Antwort: {antwort[:100]}...")

                # Chat-History aktualisieren
                if antwort:  # Check if antwort is defined
                    chat_history.append({"user": user_message, "bot": antwort})
                    session[chat_key] = chat_history
                    store_chatlog(user_name, chat_history)

                    return (
                        jsonify({"response": antwort}),
                        200,
                    ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else redirect(
                        url_for("chat")
                    )
                else:
                    error_message = "Keine Antwort erhalten."
                    debug_print("API Calls", error_message)
                    return (
                        jsonify({"error": error_message}),
                        500,
                    ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else redirect(
                        url_for("chat")
                    )
                    
            except Exception as e:
                logging.exception("Fehler beim Verarbeiten der Anfrage")
                error_message = f"Fehler bei der Kommunikation: {str(e)}"
                debug_print("API Calls", error_message)
                flash("Es gab ein Problem bei der Kommunikation mit dem Bot.", "danger")
                return (
                    jsonify({"error": error_message}),
                    500,
                ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else redirect(
                    url_for("chat")
                )

        stats = calculate_chat_stats()
        
        # Prüfen, ob eine Human-in-Loop Rückfrage angezeigt werden soll
        human_in_loop_data = session.get("human_in_loop_data")
        
        # Prüfen, ob wir eine direkte Antwort haben, die angezeigt werden soll
        last_response = None
        if session.get("last_response"):
            last_response = session.pop("last_response")
            # Wenn wir eine direkte Antwort haben, fügen wir sie der Chat-History hinzu
            if last_response and chat_key in session:
                # Wir haben die Benutzeranfrage schon beim Speichern der Antwort hinzugefügt
                # Jetzt müssen wir nur sicherstellen, dass die Antwort in der History ist
                session.modified = True
        
        # Wenn human_in_loop_data existiert, geben wir die Rückfrage-Optionen mit an das Template
        return render_template("chat.html", 
                              chat_history=display_chat_history, 
                              stats=stats,
                              human_in_loop=human_in_loop_data,
                              last_response=last_response)

    except Exception as e:
        logging.exception("Fehler in chat-Funktion.")
        flash("Ein unerwarteter Fehler ist aufgetreten.", "danger")
        return (
            jsonify({"error": "Interner Serverfehler."}),
            500,
        ) if request.headers.get("X-Requested-With") == "XMLHttpRequest" else "Interner Serverfehler", 500

@app.route('/test_bigquery')
def test_bigquery():
    try:
        service_account_path = '/home/PfS/gcpxbixpflegehilfesenioren-a47c654480a8.json'
        client = bigquery.Client.from_service_account_json(service_account_path)
        
        # E-Mail aus Session holen
        email = session.get('email') or session.get('google_user_email', '')
        
        # Abfrage mit WHERE-Klausel für User-Info
        user_query = """
        SELECT email, _id 
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.proto_users` 
        WHERE email = @email
        LIMIT 10
        """
        
        # Parameter definieren für User-Query
        user_job_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("email", "STRING", email)
            ]
        )
        
        # User-Abfrage ausführen
        user_query_job = client.query(user_query, job_config=user_job_config)
        user_results = user_query_job.result()
        
        # HTML-Ausgabe mit Bootstrap für besseres Layout
        output = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>BigQuery Dashboard zum Testen der geladenen Daten</title>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1">
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
            <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
            <style>
                .table-responsive { margin-bottom: 20px; }
                .tab-pane { padding: 15px; }
            </style>
        </head>
        <body>
        <div class="container mt-4">
            <h1>BigQuery Dashboard zum Testen der geladenen Daten</h1>
            <p>Abfrage für E-Mail: """ + email + """</p>
        """
        
        # Seller-ID ermitteln
        seller_id = session.get('seller_id')
        rows_found = False
        
        output += """
        <h2>Benutzerinformationen</h2>
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead><tr><th>E-Mail</th><th>ID</th></tr></thead>
            <tbody>
        """
        
        for row in user_results:
            rows_found = True
            output += f"<tr><td>{row['email']}</td><td>{row['_id']}</td></tr>"
            if not seller_id:  # Wenn keine seller_id in der Session ist, nutze die aus der Abfrage
                seller_id = row['_id']
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not rows_found:
            output += f"<p class='alert alert-warning'>Keine Daten für E-Mail '{email}' gefunden!</p>"
            
        # Wenn wir keine seller_id haben, können wir keine weiteren Abfragen durchführen
        if not seller_id:
            output += "<p class='alert alert-danger'>Keine Seller ID gefunden für weitere Abfragen!</p></div></body></html>"
            return output
            
        output += f"<p>Verwende Seller ID: <strong>{seller_id}</strong> für weitere Abfragen</p>"
        
        # Tabs für verschiedene Datenquellen
        output += """
        <ul class="nav nav-tabs" id="dataTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="carestays-tab" data-bs-toggle="tab" data-bs-target="#carestays" type="button" role="tab">Care Stays</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="contracts-tab" data-bs-toggle="tab" data-bs-target="#contracts" type="button" role="tab">Verträge</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="leads-tab" data-bs-toggle="tab" data-bs-target="#leads" type="button" role="tab">Leads</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="caregivers-tab" data-bs-toggle="tab" data-bs-target="#caregivers" type="button" role="tab">Pflegekräfte</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="tickets-tab" data-bs-toggle="tab" data-bs-target="#tickets" type="button" role="tab">Tickets</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="agencies-tab" data-bs-toggle="tab" data-bs-target="#agencies" type="button" role="tab">Agenturen</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="stats-tab" data-bs-toggle="tab" data-bs-target="#stats" type="button" role="tab">Statistiken</button>
            </li>
        </ul>
        
        <div class="tab-content" id="dataTabsContent">
        """
        
        # ==================== CARE STAYS TAB ====================
        output += """<div class="tab-pane fade show active" id="carestays" role="tabpanel">
            <h3 class="mt-3">Aktive Care Stays</h3>"""
        
        # Care Stays Abfrage
        care_stays_query = """
        SELECT
            cs.bill_start,
            cs.bill_end,
            cs.arrival,
            cs.departure,
            cs.presented_at,
            cs.contract_id,
            cs.created_at,
            cs.stage,
            cs.prov_seller AS seller_prov,
            cs._id AS cs_id,
            cs.care_giver_instance_id,
            TIMESTAMP(cs.bill_end) AS parsed_bill_end,
            TIMESTAMP(cs.bill_start) AS parsed_bill_start,
            DATE_DIFF(
                DATE(TIMESTAMP(cs.bill_end)),
                DATE(TIMESTAMP(cs.bill_start)),
                DAY
            ) AS care_stay_duration_days,
            c.agency_id,
            c.active,
            c.termination_reason,
            l._id AS lead_id,
            l.tracks AS lead_tracks,
            l.created_at AS lead_created_at,
            lead_names.first_name,
            lead_names.last_name,
            agencies.name AS agency_name
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_stays` AS cs
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` AS c ON cs.contract_id = c._id
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` AS h ON c.household_id = h._id
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` AS l ON h.lead_id = l._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_names 
            ON l._id = lead_names._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.agencies` AS agencies
            ON c.agency_id = agencies._id
        WHERE l.seller_id = @seller_id
          AND cs.stage = 'Bestätigt'
          AND DATE(TIMESTAMP(cs.bill_end)) >= CURRENT_DATE()
          AND cs.rejection_reason IS NULL
        ORDER BY cs.bill_start DESC
        LIMIT 100
        """
        
        care_stays_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        # Care Stays-Abfrage ausführen
        care_stays_job = client.query(care_stays_query, job_config=care_stays_config)
        care_stays_results = care_stays_job.result()
        
        # Tabelle für Care Stays ausgeben
        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>CS ID</th>
                    <th>Lead Name</th>
                    <th>Agency</th>
                    <th>Bill Start</th>
                    <th>Bill End</th>
                    <th>Arrival</th>
                    <th>Departure</th>
                    <th>Status</th>
                    <th>Prov. Verkäufer</th>
                    <th>Dauer (Tage)</th>
                </tr>
            </thead>
            <tbody>
        """
        
        care_stays_found = False
        for row in care_stays_results:
            care_stays_found = True
            lead_id = row['lead_id']
            
            # Zusammensetzen des Kundennamens
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            lead_name = f"{first_name} {last_name}".strip() if (first_name or last_name) else f"Lead ID: {lead_id}"
            
            # Agency Name
            agency_name = row.get('agency_name', 'N/A')
            
            # Sichere Anzeige der Datumswerte
            bill_start = str(row['bill_start']) if row['bill_start'] else 'N/A'
            bill_end = str(row['bill_end']) if row['bill_end'] else 'N/A'
            arrival = str(row['arrival']) if row.get('arrival') else 'N/A'
            departure = str(row['departure']) if row.get('departure') else 'N/A'
            
            output += f"""
            <tr>
                <td>{row['cs_id']}</td>
                <td>{lead_name}</td>
                <td>{agency_name}</td>
                <td>{bill_start}</td>
                <td>{bill_end}</td>
                <td>{arrival}</td>
                <td>{departure}</td>
                <td>{row['stage']}</td>
                <td>{row['seller_prov']}</td>
                <td>{row['care_stay_duration_days']}</td>
            </tr>"""
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not care_stays_found:
            output += f"<p class='alert alert-warning'>Keine aktiven Care Stays für Seller ID '{seller_id}' gefunden!</p>"
        
        output += "</div>"  # Ende des Care Stays Tab
        
        # ==================== VERTRÄGE TAB ====================
        output += """<div class="tab-pane fade" id="contracts" role="tabpanel">
            <h3 class="mt-3">Verträge</h3>"""
        
        contracts_query = """
        SELECT
            c._id AS contract_id,
            c.active,
            c.termination_reason,
            c.agency_id,
            c.household_id,
            h.lead_id,
            agencies.name AS agency_name,
            lead_names.first_name,
            lead_names.last_name,
            c.created_at
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` AS c
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` AS h ON c.household_id = h._id
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` AS l ON h.lead_id = l._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_names 
            ON l._id = lead_names._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.agencies` AS agencies
            ON c.agency_id = agencies._id
        WHERE l.seller_id = @seller_id
        ORDER BY c.created_at DESC
        LIMIT 100
        """
        
        contracts_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        contracts_job = client.query(contracts_query, job_config=contracts_config)
        contracts_results = contracts_job.result()
        
        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>Contract ID</th>
                    <th>Lead</th>
                    <th>Agentur</th>
                    <th>Aktiv</th>
                    <th>Kündigungsgrund</th>
                    <th>Erstellt am</th>
                </tr>
            </thead>
            <tbody>
        """
        
        contracts_found = False
        for row in contracts_results:
            contracts_found = True
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            lead_name = f"{first_name} {last_name}".strip() if (first_name or last_name) else f"Lead ID: {row['lead_id']}"
            
            agency_name = row.get('agency_name', 'N/A')
            created_at = str(row['created_at']) if row['created_at'] else 'N/A'
            active = "Ja" if row['active'] else "Nein"
            
            output += f"""
            <tr>
                <td>{row['contract_id']}</td>
                <td>{lead_name}</td>
                <td>{agency_name}</td>
                <td>{active}</td>
                <td>{row['termination_reason'] or 'N/A'}</td>
                <td>{created_at}</td>
            </tr>"""
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not contracts_found:
            output += f"<p class='alert alert-warning'>Keine Verträge für Seller ID '{seller_id}' gefunden!</p>"
        
        output += "</div>"  # Ende des Verträge Tab
        
        # ==================== LEADS TAB ====================
        output += """<div class="tab-pane fade" id="leads" role="tabpanel">
            <h3 class="mt-3">Leads</h3>"""
        
        leads_query = """
        SELECT
            l._id AS lead_id,
            la.first_name,
            la.last_name,
            l.created_at AS lead_created_at,
            l.source_data
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` AS l
        JOIN `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS la
        ON la._id = l._id
        WHERE l.seller_id = @seller_id
        ORDER BY l.created_at DESC
        LIMIT 100
        """
        
        leads_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        leads_job = client.query(leads_query, job_config=leads_config)
        leads_results = leads_job.result()
        
        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>Lead ID</th>
                    <th>Vorname</th>
                    <th>Nachname</th>
                    <th>Erstellt am</th>
                </tr>
            </thead>
            <tbody>
        """
        
        leads_found = False
        for row in leads_results:
            leads_found = True
            created_at = str(row['lead_created_at']) if row['lead_created_at'] else 'N/A'
            
            output += f"""
            <tr>
                <td>{row['lead_id']}</td>
                <td>{row.get('first_name', 'N/A')}</td>
                <td>{row.get('last_name', 'N/A')}</td>
                <td>{created_at}</td>
            </tr>"""
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not leads_found:
            output += f"<p class='alert alert-warning'>Keine Leads für Seller ID '{seller_id}' gefunden!</p>"
        
        output += "</div>"  # Ende des Leads Tab
        
        # ==================== PFLEGEKRÄFTE TAB ====================
        output += """<div class="tab-pane fade" id="caregivers" role="tabpanel">
            <h3 class="mt-3">Pflegekräfte</h3>"""
        
        caregivers_query = """
        SELECT
            cgi._id AS care_giver_instance_id,
            cg.first_name AS giver_first_name,
            cg.last_name AS giver_last_name,
            cs._id AS carestay_id,
            c._id AS contract_id,
            h.lead_id,
            lead_names.first_name AS lead_first_name,
            lead_names.last_name AS lead_last_name
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_giver_instances` cgi
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_givers` cg ON cgi.care_giver_id = cg._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_stays` cs ON cs.care_giver_instance_id = cgi._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` c ON cs.contract_id = c._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` h ON c.household_id = h._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` l ON h.lead_id = l._id
        LEFT JOIN `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_names 
            ON l._id = lead_names._id
        WHERE l.seller_id = @seller_id
        LIMIT 100
        """
        
        caregivers_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        caregivers_job = client.query(caregivers_query, job_config=caregivers_config)
        caregivers_results = caregivers_job.result()
        
        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>Pflegekraft</th>
                    <th>Kunde</th>
                    <th>Care Stay ID</th>
                    <th>Instance ID</th>
                </tr>
            </thead>
            <tbody>
        """
        
        caregivers_found = False
        for row in caregivers_results:
            caregivers_found = True
            caregiver_name = f"{row.get('giver_first_name', '')} {row.get('giver_last_name', '')}".strip() or 'N/A'
            lead_name = f"{row.get('lead_first_name', '')} {row.get('lead_last_name', '')}".strip() or 'N/A'
            
            output += f"""
            <tr>
                <td>{caregiver_name}</td>
                <td>{lead_name}</td>
                <td>{row.get('carestay_id', 'N/A')}</td>
                <td>{row['care_giver_instance_id']}</td>
            </tr>"""
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not caregivers_found:
            output += f"<p class='alert alert-warning'>Keine Pflegekräfte für Seller ID '{seller_id}' gefunden!</p>"
        
        output += "</div>"  # Ende des Pflegekräfte Tab
        
        # ==================== TICKETS TAB ====================
        output += """<div class="tab-pane fade" id="tickets" role="tabpanel">
            <h3 class="mt-3">Tickets</h3>"""

        tickets_query = """
        SELECT
            t.subject,
            t.messages,
            t.created_at,
            tc.Datum,
            t.updated_at,
            t._id,
            t.ticketable_id,
            t.ticketable_type,
            tc.seller,
            tc.agency,
            -- Lead-Informationen
            CASE 
                WHEN t.ticketable_type = 'Lead' THEN lead_direct.first_name
                WHEN t.ticketable_type = 'Contract' THEN lead_contract.first_name
                WHEN t.ticketable_type = 'CareStay' THEN lead_carestay.first_name
                WHEN t.ticketable_type = 'Visor' THEN lead_visor.first_name
                ELSE NULL
            END AS lead_first_name,
            CASE 
                WHEN t.ticketable_type = 'Lead' THEN lead_direct.last_name
                WHEN t.ticketable_type = 'Contract' THEN lead_contract.last_name
                WHEN t.ticketable_type = 'CareStay' THEN lead_carestay.last_name
                WHEN t.ticketable_type = 'Visor' THEN lead_visor.last_name
                ELSE NULL
            END AS lead_last_name,
            CASE
                WHEN t.ticketable_type = 'Lead' THEN t.ticketable_id
                WHEN t.ticketable_type = 'Contract' THEN contract_lead.lead_id
                WHEN t.ticketable_type = 'CareStay' THEN carestay_lead.lead_id
                WHEN t.ticketable_type = 'Visor' THEN visor_lead.lead_id
                ELSE NULL
            END AS lead_id
        FROM
            `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.tickets` t
        LEFT JOIN
            `gcpxbixpflegehilfesenioren.dataform_staging.tickets_creation_end` tc
            ON t._id = tc.Ticket_ID

        -- Direkter Lead-Join für Lead-Tickets
        LEFT JOIN
            `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_direct
            ON t.ticketable_type = 'Lead' AND t.ticketable_id = lead_direct._id

        -- Contract-Lead-Join für Contract-Tickets
        LEFT JOIN (
            SELECT
                c._id AS contract_id,
                h.lead_id,
                l._id AS lead_orig_id
            FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` c
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` h ON c.household_id = h._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` l ON h.lead_id = l._id
        ) AS contract_lead
            ON t.ticketable_type = 'Contract' AND t.ticketable_id = contract_lead.contract_id
        LEFT JOIN
            `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_contract
            ON contract_lead.lead_orig_id = lead_contract._id

        -- CareStay-Lead-Join für CareStay-Tickets
        LEFT JOIN (
            SELECT
                cs._id AS carestay_id,
                h.lead_id,
                l._id AS lead_orig_id
            FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_stays` cs
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` c ON cs.contract_id = c._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` h ON c.household_id = h._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` l ON h.lead_id = l._id
        ) AS carestay_lead
            ON t.ticketable_type = 'CareStay' AND t.ticketable_id = carestay_lead.carestay_id
        LEFT JOIN
            `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_carestay
            ON carestay_lead.lead_orig_id = lead_carestay._id

        -- Visor-Lead-Join für Visor-Tickets
        LEFT JOIN (
            SELECT
                v._id AS visor_id,
                h.lead_id,
                l._id AS lead_orig_id
            FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.visors` v
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.postings` p ON v.posting_id = p._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` h ON p.household_id = h._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` l ON h.lead_id = l._id
        ) AS visor_lead
            ON t.ticketable_type = 'Visor' AND t.ticketable_id = visor_lead.visor_id
        LEFT JOIN
            `gcpxbixpflegehilfesenioren.dataform_staging.leads_and_seller_and_source_with_address` AS lead_visor
            ON visor_lead.lead_orig_id = lead_visor._id

        WHERE
            tc.seller = 'Pflegeteam Heer'
        ORDER BY 
            t.created_at DESC
        LIMIT 100
        """

        tickets_job = client.query(tickets_query)
        tickets_results = tickets_job.result()

        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>Ticket ID</th>
                    <th>Betreff</th>
                    <th>Typ</th>
                    <th>Lead</th>
                    <th>Erstellt am</th>
                    <th>Aktualisiert am</th>
                    <th>Agentur</th>
                </tr>
            </thead>
            <tbody>
        """

        tickets_found = False
        for row in tickets_results:
            tickets_found = True
            created_at = str(row['created_at']) if row['created_at'] else 'N/A'
            updated_at = str(row['updated_at']) if row['updated_at'] else 'N/A'
            
            # Lead-Namen zusammensetzen
            lead_name = "N/A"
            if row.get('lead_first_name') or row.get('lead_last_name'):
                lead_name = f"{row.get('lead_first_name', '')} {row.get('lead_last_name', '')}".strip()
            
            output += f"""
            <tr>
                <td>{row['_id']}</td>
                <td>{row['subject']}</td>
                <td>{row['ticketable_type'] or 'N/A'}</td>
                <td>{lead_name}</td>
                <td>{created_at}</td>
                <td>{updated_at}</td>
                <td>{row['agency'] or 'N/A'}</td>
            </tr>"""

        output += """
            </tbody>
        </table>
        </div>
        """

        if not tickets_found:
            output += f"<p class='alert alert-warning'>Keine Tickets für '{seller_id}' gefunden!</p>"

        output += "</div>"  # Ende des Tickets Tab
        
        # ==================== AGENTUREN TAB ====================
        output += """<div class="tab-pane fade" id="agencies" role="tabpanel">
            <h3 class="mt-3">Agenturen</h3>"""
        
        agencies_query = """
        SELECT
            _id AS agency_id,
            name AS agency_name
        FROM gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.agencies
        ORDER BY name
        LIMIT 100
        """
        
        agencies_job = client.query(agencies_query)
        agencies_results = agencies_job.result()
        
        output += """
        <div class="table-responsive">
        <table class="table table-striped table-bordered">
            <thead>
                <tr>
                    <th>Agentur ID</th>
                    <th>Name</th>
                </tr>
            </thead>
            <tbody>
        """
        
        agencies_found = False
        for row in agencies_results:
            agencies_found = True
            output += f"""
            <tr>
                <td>{row['agency_id']}</td>
                <td>{row['agency_name']}</td>
            </tr>"""
        
        output += """
            </tbody>
        </table>
        </div>
        """
        
        if not agencies_found:
            output += f"<p class='alert alert-warning'>Keine Agenturen gefunden!</p>"
        
        output += "</div>"  # Ende des Agenturen Tab
        
        # ==================== STATISTIKEN TAB ====================
        output += """<div class="tab-pane fade" id="stats" role="tabpanel">
            <h3 class="mt-3">Statistiken</h3>"""
        
        # Care Stay Statistiken
        stats_query = """
        SELECT
            COUNT(DISTINCT cs._id) AS total_care_stays,
            COUNT(DISTINCT c._id) AS total_contracts,
            COUNT(DISTINCT l._id) AS total_leads,
            AVG(DATE_DIFF(
                DATE(TIMESTAMP(cs.bill_end)),
                DATE(TIMESTAMP(cs.bill_start)),
                DAY
            )) AS avg_care_stay_duration,
            SUM(CAST(cs.prov_seller AS FLOAT64)) AS total_prov_seller
        FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_stays` AS cs
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` AS c ON cs.contract_id = c._id
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` AS h ON c.household_id = h._id
        JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` AS l ON h.lead_id = l._id
        WHERE l.seller_id = @seller_id
          AND cs.stage = 'Bestätigt'
        """
        
        stats_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        stats_job = client.query(stats_query, job_config=stats_config)
        stats_results = stats_job.result()
        
        # Monatliche Statistiken
        monthly_stats_query = """
        WITH monthly_data AS (
            SELECT
                FORMAT_DATE('%Y-%m', DATE(TIMESTAMP(cs.bill_start))) AS month,
                COUNT(DISTINCT cs._id) AS new_care_stays,
                SUM(CAST(cs.prov_seller AS FLOAT64)) AS monthly_prov
            FROM `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.care_stays` AS cs
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.contracts` AS c ON cs.contract_id = c._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.households` AS h ON c.household_id = h._id
            JOIN `gcpxbixpflegehilfesenioren.PflegehilfeSeniore_BI.leads` AS l ON h.lead_id = l._id
            WHERE l.seller_id = @seller_id
              AND cs.stage = 'Bestätigt'
              AND DATE(TIMESTAMP(cs.bill_start)) >= DATE_SUB(CURRENT_DATE(), INTERVAL 12 MONTH)
            GROUP BY month
            ORDER BY month DESC
        )
        SELECT * FROM monthly_data
        """
        
        monthly_stats_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("seller_id", "STRING", seller_id)
            ]
        )
        
        monthly_stats_job = client.query(monthly_stats_query, job_config=monthly_stats_config)
        monthly_stats_results = monthly_stats_job.result()
        
        # Gesamtstatistiken
        output += """<div class="row">
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">Gesamtstatistiken</div>
                    <div class="card-body">
                        <table class="table">
                            <thead>
                                <tr>
                                    <th>Metrik</th>
                                    <th>Wert</th>
                                </tr>
                            </thead>
                            <tbody>
        """
        
        for row in stats_results:
            output += f"""
                <tr><td>Gesamtzahl Care Stays</td><td>{row['total_care_stays']}</td></tr>
                <tr><td>Gesamtzahl Verträge</td><td>{row['total_contracts']}</td></tr>
                <tr><td>Gesamtzahl Leads</td><td>{row['total_leads']}</td></tr>
                <tr><td>Durchschnittliche Care Stay Dauer (Tage)</td><td>{row['avg_care_stay_duration']:.2f}</td></tr>
                <tr><td>Gesamtprovision</td><td>{row['total_prov_seller']:.2f} €</td></tr>
            """
        
        output += """
                            </tbody>
                        </table>
                    </div>
                </div>
            </div>
            
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">Monatliche Statistiken</div>
                    <div class="card-body">
                        <table class="table">
                            <thead>
                                <tr>
                                    <th>Monat</th>
                                    <th>Neue Care Stays</th>
                                    <th>Provision</th>
                                </tr>
                            </thead>
                            <tbody>
        """
        
        monthly_stats_found = False
        for row in monthly_stats_results:
            monthly_stats_found = True
            output += f"""
                <tr>
                    <td>{row['month']}</td>
                    <td>{row['new_care_stays']}</td>
                    <td>{row['monthly_prov']:.2f} €</td>
                </tr>
            """
        
        if not monthly_stats_found:
            output += "<tr><td colspan='3'>Keine monatlichen Daten gefunden</td></tr>"
        
        output += """
                            </tbody>
                        </table>
                    </div>
                </div>
            </div>
        </div>"""
        
        output += "</div>"  # Ende des Statistiken Tab
        
        # Abschluss der HTML-Struktur
        output += """
        </div>  <!-- Ende der Tab-Inhalte -->
        </div>  <!-- Ende des Containers -->
        </body>
        </html>
        """
        
        return output
    except Exception as e:
        error_output = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>BigQuery Test - Fehler</title>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1">
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
        </head>
        <body>
        <div class="container mt-4">
            <div class="alert alert-danger">
                <h4>Fehler beim Ausführen der BigQuery-Abfragen:</h4>
                <p>{str(e)}</p>
                <pre>{traceback.format_exc()}</pre>
            </div>
        </div>
        </body>
        </html>
        """
        return error_output

@app.route('/get_active_care_stays_now', methods=['GET'])
def get_active_care_stays_now():
    """
    Liefert die aktiven Care Stays für den Verkäufer als JSON
    """
    try:
        # Seller ID aus der Session holen
        seller_id = session.get('seller_id')
        if not seller_id:
            return jsonify({
                "error": "Keine Seller ID in der Session gefunden",
                "status": "error"
            }), 401
        
        # Lade die Abfragemuster
        with open('query_patterns.json', 'r', encoding='utf-8') as f:
            query_patterns = json.load(f)
        
        # Hole die "get_active_care_stays_now" Abfrage
        query_name = "get_active_care_stays_now"
        if query_name not in query_patterns['common_queries']:
            return jsonify({
                "error": f"Abfrage {query_name} nicht gefunden",
                "status": "error"
            }), 404
        
        query_pattern = query_patterns['common_queries'][query_name]
        
        # Parameter für die Abfrage vorbereiten
        parameters = {'seller_id': seller_id, 'limit': 100}
        
        # Führe die Abfrage aus
        result = execute_bigquery_query(
            query_pattern['sql_template'],
            parameters
        )
        
        # Formatiere das Ergebnis
        formatted_result = format_query_result(result, query_pattern.get('result_structure'))
        
        return jsonify({
            "data": formatted_result,
            "count": len(formatted_result),
            "status": "success"
        })
    
    except Exception as e:
        error_trace = traceback.format_exc()
        logging.error(f"Fehler in get_active_care_stays_now: {str(e)}\n{error_trace}")
        return jsonify({
            "error": str(e),
            "trace": error_trace,
            "status": "error"
        }), 500

@app.route('/get_dashboard_data', methods=['GET'])
def get_dashboard_data():
    """
    Liefert Daten für das Dashboard, basierend auf dem angegebenen Abfragetyp.
    Diese Route wird beim Öffnen des Dashboards aufgerufen.
    """
    try:
        # Seller ID aus der Session holen
        seller_id = session.get('seller_id')
        logging.info(f"Dashboard: Seller ID aus Session: {seller_id}")
        
        if not seller_id:
            logging.error("Dashboard: Keine Seller ID in der Session gefunden")
            return jsonify({
                "error": "Keine Seller ID in der Session gefunden",
                "status": "error"
            }), 401
        
        # Lade die Abfragemuster
        logging.info("Dashboard: Lade Abfragemuster")
        with open('query_patterns.json', 'r', encoding='utf-8') as f:
            query_patterns = json.load(f)
        
        # Dashboard-Daten sammeln
        dashboard_result = {}
        
        # Hole die "get_active_care_stays_now" Abfrage
                # 1. Aktive Kunden (get_active_care_stays_now)
        query_name = "get_active_care_stays_now"
        logging.info(f"Dashboard: Verwende Abfrage {query_name}")
        
        if query_name not in query_patterns['common_queries']:
            logging.error(f"Dashboard: Abfrage {query_name} nicht gefunden")
            return jsonify({
                "error": f"Abfrage {query_name} nicht gefunden",
                "status": "error"
            }), 404
        
        query_pattern = query_patterns['common_queries'][query_name]
        
        # Parameter dynamisch setzen
        parameters = {
            'seller_id': seller_id,
            'start_of_month': datetime.now().date().replace(day=1).isoformat(),
            'end_of_month': datetime.now().date().isoformat(),
            'days_in_month': calendar.monthrange(datetime.now().year, datetime.now().month)[1]
        }
        logging.info(f"Dashboard: Parameter: {parameters}")

        result = execute_bigquery_query(
            query_pattern['sql_template'],
            parameters
        )
        
        # Formatiere das Ergebnis
        formatted_result = format_query_result(result, query_pattern.get('result_structure'))
        logging.info(f"Dashboard: Ergebnis formatiert, {len(formatted_result)} Einträge")
        
        # Speichern für die Antwort
        dashboard_result['active_customers'] = {
            "data": formatted_result,
            "count": len(formatted_result)
        }


                # 2. Abschlussquote (get_cvr_lead_contract)
        query_name = "get_cvr_lead_contract"
        logging.info(f"Dashboard: Verwende Abfrage {query_name}")
        
        if query_name in query_patterns['common_queries']:
            query_pattern = query_patterns['common_queries'][query_name]
            
            # Zeitraum: Letzte 30 Tage
            end_date = datetime.now().date().isoformat()
            start_date = (datetime.now().date() - timedelta(days=30)).isoformat()
            parameters = {'seller_id': seller_id, 'start_date': start_date, 'end_date': end_date}
            logging.info(f"Dashboard: Parameter für Abschlussquote: {parameters}")
            
            # Führe die Abfrage aus
            logging.info("Dashboard: Führe BigQuery-Abfrage für Abschlussquote aus")
            cvr_result = execute_bigquery_query(
                query_pattern['sql_template'],
                parameters
            )
            
            # Formatiere das Ergebnis
            formatted_cvr = format_query_result(cvr_result, query_pattern.get('result_structure'))
            logging.info(f"Dashboard: Abschlussquotenabfrage abgeschlossen")
            
            # Speichern für die Antwort
            dashboard_result['conversion_rate'] = formatted_cvr[0] if formatted_cvr else {}
        else:
            logging.error(f"Dashboard: Abfrage {query_name} nicht gefunden")
            dashboard_result['conversion_rate'] = {}

                # 3. Neue Verträge der letzten 14 Tage (get_contract_count)
        query_name = "get_contract_count"
        logging.info(f"Dashboard: Verwende Abfrage {query_name}")
        
        if query_name in query_patterns['common_queries']:
            query_pattern = query_patterns['common_queries'][query_name]
            
            # Zeitraum: Letzte 14 Tage
            end_date = datetime.now().date().isoformat()
            start_date = (datetime.now().date() - timedelta(days=14)).isoformat()
            parameters = {'seller_id': seller_id, 'start_date': start_date, 'end_date': end_date}
            logging.info(f"Dashboard: Parameter für Neue Verträge: {parameters}")
            
            # Führe die Abfrage aus
            logging.info("Dashboard: Führe BigQuery-Abfrage für Neue Verträge aus")
            contracts_result = execute_bigquery_query(
                query_pattern['sql_template'],
                parameters
            )
            
            # Formatiere das Ergebnis
            formatted_contracts = format_query_result(contracts_result, query_pattern.get('result_structure'))
            logging.info(f"Dashboard: Neue Verträge Abfrage abgeschlossen")
            
            # Speichern für die Antwort
            dashboard_result['new_contracts'] = formatted_contracts[0] if formatted_contracts else {}
        else:
            logging.error(f"Dashboard: Abfrage {query_name} nicht gefunden")
            dashboard_result['new_contracts'] = {}

                # 4. Kündigungen (get_contract_terminations)
        query_name = "get_contract_terminations"
        logging.info(f"Dashboard: Verwende Abfrage {query_name}")
        
        if query_name in query_patterns['common_queries']:
            query_pattern = query_patterns['common_queries'][query_name]
            
            # Zeitraum: Letzte 30 Tage
            end_date = datetime.now().date().isoformat()
            start_date = (datetime.now().date() - timedelta(days=30)).isoformat()
            parameters = {'seller_id': seller_id, 'start_date': start_date, 'end_date': end_date, 'limit': 500}
            logging.info(f"Dashboard: Parameter für Kündigungen: {parameters}")
            
            # Führe die Abfrage aus
            logging.info("Dashboard: Führe BigQuery-Abfrage für Kündigungen aus")
            terminations_result = execute_bigquery_query(
                query_pattern['sql_template'],
                parameters
            )
            
            # Formatiere das Ergebnis
            formatted_terminations = format_query_result(terminations_result, query_pattern.get('result_structure'))
            logging.info(f"Dashboard: Kündigungen Abfrage abgeschlossen")
            
            # Extrahiere die relevanten Kennzahlen - bei der ersten Zeile stehen die Gesamtzahlen
            if formatted_terminations and len(formatted_terminations) > 0:
                terminations_data = {
                    'serious_terminations_count': formatted_terminations[0].get('serious_terminations_count', 0),
                    'agency_switch_count': formatted_terminations[0].get('agency_switch_count', 0),
                    'total_terminations_count': formatted_terminations[0].get('total_terminations_count', 0)
                }
            else:
                terminations_data = {
                    'serious_terminations_count': 0,
                    'agency_switch_count': 0,
                    'total_terminations_count': 0
                }
            
            # Speichern für die Antwort
            dashboard_result['terminations'] = terminations_data
        else:
            logging.error(f"Dashboard: Abfrage {query_name} nicht gefunden")
            dashboard_result['terminations'] = {'serious_terminations_count': 0, 'agency_switch_count': 0, 'total_terminations_count': 0}

                # 5. Umsatz Pro Rata (laufender Monat)
        query_name_revenue = "get_revenue_current_month_pro_rata"
        logging.info(f"Dashboard: Verwende Abfrage {query_name_revenue}")
        if query_name_revenue in query_patterns['common_queries']:
            query_pattern_revenue = query_patterns['common_queries'][query_name_revenue]
            parameters_revenue = {
                'seller_id': seller_id,
                'start_of_month': datetime.now().date().replace(day=1).isoformat(),
                'end_of_month': datetime.now().date().isoformat(),
                'days_in_month': calendar.monthrange(datetime.now().year, datetime.now().month)[1]
            }
            logging.info(f"Dashboard: Parameter für Umsatz Pro Rata: {parameters_revenue}")

            revenue_result = execute_bigquery_query(
                query_pattern_revenue['sql_template'],
                parameters_revenue
            )
            formatted_revenue = format_query_result(revenue_result, query_pattern_revenue.get('result_structure'))
            logging.info(f"Dashboard: Umsatz Pro Rata Abfrage abgeschlossen")

            # Nimm den Wert oder 0, falls kein Ergebnis
            dashboard_result['pro_rata_revenue'] = formatted_revenue[0]['total_monthly_pro_rata_revenue'] if formatted_revenue and formatted_revenue[0] else 0
        else:
            logging.error(f"Dashboard: Abfrage {query_name_revenue} nicht gefunden")
            dashboard_result['pro_rata_revenue'] = 0

                # Gesamte Antwort zusammenstellen
        response = {
            "data": dashboard_result['active_customers']['data'],
            "count": dashboard_result['active_customers']['count'],
            "conversion_rate": dashboard_result['conversion_rate'],
            "new_contracts": dashboard_result['new_contracts'],
            "terminations": dashboard_result['terminations'],
            "pro_rata_revenue": dashboard_result.get('pro_rata_revenue', 0), # Hier hinzugefügt
            "status": "success"
        }
        logging.info("Dashboard: Sende Antwort")
        return jsonify(response)
    
    except Exception as e:
        error_trace = traceback.format_exc()
        logging.error(f"Fehler in get_dashboard_data: {str(e)}\n{error_trace}")
        return jsonify({
            "error": str(e),
            "trace": error_trace,
            "status": "error"
        }), 500

@app.route('/update_stream_chat_history', methods=['POST'])
def update_stream_chat_history():
    """Update chat history in the session from streaming responses"""
    try:
        data = request.json
        user_message = data.get('user_message')
        bot_response = data.get('bot_response')
        
        if not user_message or not bot_response:
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400
            
        user_id = session.get("user_id")
        user_name = session.get("user_name")
        
        if not user_id or not user_name:
            return jsonify({'success': False, 'error': 'No active session'}), 400
            
        chat_key = f'chat_history_{user_id}'
        
        # Get current chat history or initialize it
        chat_history = session.get(chat_key, [])
        
        # Add the new messages
        chat_history.append({"user": user_message, "bot": bot_response})
        
        # Update the session
        session[chat_key] = chat_history
        session.modified = True
        
        # Store in the persistent storage
        store_chatlog(user_name, chat_history)
        
        return jsonify({'success': True})
        
    except Exception as e:
        logging.exception("Error updating chat history")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/check_login')
def check_login():
    # Alle relevanten Session-Daten anzeigen
    session_data = {
        'user_id': session.get('user_id'),
        'user_name': session.get('user_name'),
        'email': session.get('email'),
        'google_user_email': session.get('google_user_email'),
        'seller_id': session.get('seller_id'),
        'is_logged_via_google': session.get('is_logged_via_google')
    }
    
    output = "<h1>Login-Status</h1>"
    output += "<pre>" + json.dumps(session_data, indent=2) + "</pre>"
    
    # Wenn E-Mail vorhanden ist, teste BigQuery-Abfrage
    email = session.get('email') or session.get('google_user_email')
    if email:
        output += f"<h2>Test der seller_id-Abfrage für {email}</h2>"
        try:
            seller_id = get_user_id_from_email(email)
            output += f"<p>Gefundene seller_id: {seller_id}</p>"
        except Exception as e:
            output += f"<p style='color:red'>Fehler: {str(e)}</p>"
    
    # Link zum Zurücksetzen der Session
    output += "<p><a href='/reset_session'>Session zurücksetzen</a></p>"
    
    return output

@app.route('/reset_session')
def reset_session():
    # Alle Session-Daten löschen, außer user_id
    user_id = session.get('user_id')
    session.clear()
    session['user_id'] = user_id
    session.modified = True
    return redirect('/check_login')

@app.route('/debug_dashboard', methods=['GET'])
def debug_dashboard():
    """
    Debug-Route für das Dashboard, die alle relevanten Informationen zurückgibt
    und in der Session gespeicherte Daten anzeigt.
    """
    debug_info = {
        "session_data": {
            "seller_id": session.get('seller_id'),
            "user_name": session.get('user_name'),
            "all_session_keys": list(session.keys())
        },
        "request_info": {
            "method": request.method,
            "headers": dict(request.headers),
            "endpoint": request.endpoint,
        }
    }
    
    # Lade die Abfragemuster und teste den Zugriff
    try:
        with open('query_patterns.json', 'r', encoding='utf-8') as f:
            debug_info["query_patterns_loaded"] = True
            query_patterns = json.load(f)
            
            # Prüfe, ob die benötigte Abfrage vorhanden ist
            query_name = "get_active_care_stays_now"
            if query_name in query_patterns.get('common_queries', {}):
                debug_info["query_exists"] = True
                debug_info["query_name"] = query_name
                debug_info["query_desc"] = query_patterns['common_queries'][query_name].get('description')
            else:
                debug_info["query_exists"] = False
    except Exception as e:
        debug_info["query_patterns_loaded"] = False
        debug_info["query_patterns_error"] = str(e)
    
    # HTML-Ausgabe für leichtere Lesbarkeit
    html_output = "<h1>Dashboard Debug-Informationen</h1>"
    html_output += "<pre>" + json.dumps(debug_info, indent=4) + "</pre>"
    
    # Button zum Testen des eigentlichen Endpunkts
    html_output += """
    <script>
    function testEndpoint() {
        fetch('/get_dashboard_data')
            .then(response => response.json())
            .then(data => {
                document.getElementById('result').textContent = JSON.stringify(data, null, 2);
            })
            .catch(error => {
                document.getElementById('result').textContent = 'Fehler: ' + error;
            });
    }
    </script>
    
    <button onclick="testEndpoint()">Test /get_dashboard_data</button>
    <pre id="result">Klicken Sie auf den Button, um den Endpunkt zu testen...</pre>
    """
    
    return html_output

@app.route('/toggle_notfall_mode', methods=['POST'])
def toggle_notfall_mode():
    try:
        activate = request.form.get('activate', '0')
        if activate == '1':
            session['notfall_mode'] = True
        else:
            session.pop('notfall_mode', None)
        return jsonify({
            'success': True,
            'notfall_mode_active': session.get('notfall_mode', False)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/store_feedback', methods=['POST'])
def store_feedback_route():
    try:
        data = request.get_json()
        feedback_type = data.get("feedback_type", "")
        comment = data.get("comment", "").strip()
        message = data.get("message", "").strip()  # Get the specific message being rated
        
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({"success": False, "message": "User nicht erkannt"}), 400
        
        chat_key = f'chat_history_{user_id}'
        chat_history = session.get(chat_key, [])
        
        # Call the updated store_feedback function with the message parameter
        store_feedback(feedback_type, comment, chat_history, message)
        return jsonify({"success": True}), 200
    except Exception as e:
        logging.exception("Fehler beim Speichern des Feedbacks.")
        return jsonify({"success": False, "error": str(e)}), 500


###########################################
# Neue Routen: Username setzen/auslesen
###########################################
@app.route('/get_username', methods=['GET'])
def get_username():
    try:
        user_name = session.get('user_name')
        return jsonify({'user_name': user_name}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/set_username', methods=['POST'])
def set_username():
    try:
        username = request.form.get('username', '').strip()
        if len(username) < 3:
            return jsonify({'success': False, 'message': 'Name zu kurz'}), 400
        
        session['user_name'] = username
        
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

###########################################
# Chat Route
###########################################

 


@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        # Behalten Sie die bestehende Admin-Login-Logik bei
        password = request.form.get('password', '')
        admin_password = os.getenv('ADMIN_PASSWORD', '')
        if password == admin_password:
            session['admin_logged_in'] = True
            flash('Erfolgreich als Administrator eingeloggt.', 'success')
            return redirect(url_for('admin'))
        else:
            flash('Falsches Passwort.', 'danger')
            return redirect(url_for('admin_login'))
            
    # Render admin login template
    return render_template('admin_login.html')

###########################################
# CSRF & Session Hooks
###########################################
from flask_wtf.csrf import generate_csrf

@app.context_processor
def inject_csrf_token():
    return {'csrf_token': generate_csrf()}

@app.before_request
def ensure_user_id():
    if 'user_id' not in session:
        session['user_id'] = str(uuid.uuid4())

###########################################
# Login-Decorator
###########################################

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

###########################################
# AJAX Endpoint für Human-in-the-Loop
###########################################
@app.route('/get_clarification_response', methods=['GET'])
def get_clarification_response():
    """
    AJAX endpoint to get a response after a human-in-loop clarification button was clicked.
    This allows the client to update the UI without a page refresh.
    """
    try:
        # Ensure user is authenticated
        user_id = session.get("user_id")
        if not user_id:
            return jsonify({"error": "Nicht angemeldet"}), 403
            
        # Check if we have a pending clarification response
        if "human_in_loop_clarification_response" in session:
            clarification_response = session.pop("human_in_loop_clarification_response")
            pending_query = session.get("pending_query", "")
            session.pop("pending_query", None)
            
            # Process the clarification response
            logging.info(f"Processing clarification response via AJAX: {clarification_response.get('text', '')}")
            
            # Extract the selected query and parameters
            selected_query = clarification_response.get("query")
            selected_params = clarification_response.get("params", {})
            
            # Add standard parameters
            if "seller_id" in selected_params and "seller_id" in session:
                selected_params["seller_id"] = session.get("seller_id")
            
            # Execute function call
            try:
                tool_result = handle_function_call(selected_query, selected_params)
                
                # Check for empty results
                result_data = json.loads(tool_result)
                if "data" in result_data and len(result_data["data"]) == 0:
                    # No data found, return a proper error
                    logging.warning(f"No data found for query {selected_query} with params {selected_params}")
                    return jsonify({
                        "success": False,
                        "error": f"Keine Daten für diese Anfrage gefunden."
                    }), 404
                    
            except Exception as e:
                logging.error(f"Error executing function call: {str(e)}")
                return jsonify({
                    "success": False,
                    "error": f"Fehler bei der Ausführung: {str(e)}"
                }), 500
            
            # Create enhanced system prompt for LLM response generation
            system_prompt = create_enhanced_system_prompt(selected_query)
            
            messages = [
                {"role": "developer", "content": system_prompt},
                {"role": "user", "content": pending_query},
                {"role": "function", "name": selected_query, "content": tool_result}
            ]
            
            # Generate response with LLM
            # o3-mini doesn't support function role, so we'll convert the function message to a user message
            adjusted_messages = [
                {"role": "developer", "content": system_prompt},
                {"role": "user", "content": f"User question: {pending_query}\n\nQuery result: {tool_result}"}
            ]
            
            response = openai.chat.completions.create(
                model="o3-mini",
                messages=adjusted_messages
            )
            
            final_response = response.choices[0].message.content
            
            # Update chat history
            chat_key = f"chat_history_{session.get('user_id')}"
            chat_history = session.get(chat_key, [])
            chat_history.append({"role": "user", "content": pending_query})
            chat_history.append({"role": "assistant", "content": final_response})
            session[chat_key] = chat_history
            session.modified = True
            
            # Return the response as JSON
            return jsonify({
                "success": True,
                "response": final_response
            })
        else:
            # No pending clarification response
            return jsonify({
                "success": False,
                "error": "Keine ausstehende Antwort auf Rückfrage gefunden"
            })
    except Exception as e:
        logging.exception(f"Error processing clarification response: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Fehler bei der Verarbeitung: {str(e)}"
        }), 500

###########################################
# Clear Chat
###########################################
@app.route('/clear_chat_history', methods=['POST'])
def clear_chat_history():
    try:
        user_id = session.get('user_id')
        if not user_id:
            flash('Benutzer nicht erkannt.', 'danger')
            return redirect(url_for('chat'))

        chat_key = f'chat_history_{user_id}'
        if chat_key in session:
            session.pop(chat_key)
            flash('Chatverlauf wurde erfolgreich geleert.', 'success')
        else:
            flash('Es gibt keinen Chatverlauf zu löschen.', 'info')
        return redirect(url_for('chat'))
    except Exception as e:
        logging.exception("Fehler beim Löschen Chatverlaufs.")
        flash('Ein Fehler ist aufgetreten.', 'danger')
        return redirect(url_for('chat'))

###########################################
# Login / Logout
###########################################
@app.route('/logout')
def logout():
    # Check if user is admin and handle accordingly
    if session.get('admin_logged_in'):
        session.pop('admin_logged_in', None)
        flash('Erfolgreich ausgeloggt.', 'success')
        return redirect(url_for('login'))
    
    # Clear specific user session data
    keys_to_remove = [
        'user_name', 'email', 'google_user_email', 
        'seller_id', 'is_logged_via_google', 'access_token'
    ]
    
    for key in keys_to_remove:
        if key in session:
            session.pop(key)
    
    # Completely clear Flask session to force a new Google login
    session.clear()
    
    # Setze den Session-Cookie zurück, um einen neuen Login zu erzwingen
    response = redirect(url_for('chat'))
    if 'session' in request.cookies:
        response.delete_cookie('session')
    
    flash('Erfolgreich ausgeloggt. Bitte melden Sie sich erneut an.', 'success')
    
    return response

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # Behalten Sie die bestehende Admin-Login-Logik bei
        password = request.form.get('password', '')
        admin_password = os.getenv('ADMIN_PASSWORD', '')
        if password == admin_password:
            session['admin_logged_in'] = True
            flash('Erfolgreich eingeloggt.', 'success')
            return redirect(url_for('admin'))
        else:
            flash('Falsches Passwort.', 'danger')
            return redirect(url_for('login'))
            
    # Render login template mit Google-Login-Option
    return render_template('login.html')

###########################################
# Lade Themen
###########################################
@app.route('/lade_themen', methods=['GET'], endpoint='lade_themen')
@login_required
def lade_themen_route():
    themen_dict = lade_themen()
    return jsonify(themen_dict), 200


###########################################
# Admin
###########################################
@app.route('/admin', methods=['GET', 'POST'])
@login_required
def admin():
    try:
        themen_dict = lade_themen()
        logging.debug("Themen_dict: %s", themen_dict)
        if request.method == 'POST':
            eingabe_text = request.form.get('eingabe_text', '').strip()
            if not eingabe_text:
                flash('Bitte geben Sie einen Wissenseintrag ein.', 'warning')
                return redirect(url_for('admin'))

            ki_aktiviert = (request.form.get('ki_var') == 'on')
            if ki_aktiviert:
                # KI-Extraktion (Beispiel)
                extraktion_messages = [
                    {"role": "user", "content": "Du bist ein Experte, der aus Transkripten Wissen extrahiert..."},
                    {"role": "user", "content": f"Extrahiere ohne Verluste:\n'''{eingabe_text}'''"}
                ]
                extraktion_response = contact_openai(extraktion_messages, model="o3-mini")
                if not extraktion_response:
                    flash("Fehler bei der Wissensextraktion durch die KI.", 'danger')
                    return redirect(url_for('admin'))

                try:
                    extraktion_text = extraktion_response
                    debug_print("Bearbeiten von Einträgen", "Extraktionsergebnis: " + extraktion_text)
                except Exception as e:
                    debug_print("Bearbeiten von Einträgen", f"Fehler: {e}")
                    flash("Fehler bei der Wissensextraktion.", 'danger')
                    return redirect(url_for('admin'))

                # Kategorisierung (Beispiel)
                themen_hierarchie = ''
                if os.path.exists(themen_datei):
                    with open(themen_datei, 'r', encoding='utf-8') as f:
                        themen_hierarchie = f.read()

                kategorisierung_messages = [
                    {"role": "user",
                     "content": "Du bist ein Assistent, der Texte in vorgegebene Themen..."},
                    {"role": "user",
                     "content": f"Hier die Themenhierarchie:\n\n{themen_hierarchie}\n\nText:\n{extraktion_text}"}
                ]
                kategorisierung_response = contact_openai(kategorisierung_messages, model="o3-mini")
                if not kategorisierung_response:
                    flash("Fehler bei der Themenkategorisierung.", 'danger')
                    return redirect(url_for('admin'))

                try:
                    kategorisierung_text = kategorisierung_response
                    json_match = re.search(r'\[\s*{.*}\s*\]', kategorisierung_text, re.DOTALL)
                    if json_match:
                        json_text = json_match.group(0)
                        json_text = re.sub(r'(\d+[a-z]*)\)\)', r'\1)', json_text)
                        kategorisierte_eintraege = json.loads(json_text)
                    else:
                        raise ValueError("Kein gültiger JSON-Inhalt gefunden")

                    for eintrag_kat in kategorisierte_eintraege:
                        thema = eintrag_kat.get('thema')
                        unterthema = eintrag_kat.get('unterthema')
                        beschreibung = eintrag_kat.get('beschreibung', '')
                        inhalt = eintrag_kat.get('inhalt')
                        if thema and unterthema and inhalt:
                            speichere_wissensbasis({
                                "thema": thema,
                                "unterthema": unterthema,
                                "beschreibung": beschreibung,
                                "inhalt": inhalt
                            })
                    flash("Alle Einträge erfolgreich gespeichert.", 'success')
                except Exception as e:
                    debug_print("Bearbeiten von Einträgen", f"Parsing-Fehler: {e}")
                    flash("KI-Antwort konnte nicht geparst werden.", 'danger')
            else:
                # Manuell
                ausgewähltes_thema = request.form.get('thema_var')
                ausgewähltes_unterthema = request.form.get('unterthema_var')
                beschreibung_text = request.form.get('beschreibung_text', '').strip()

                if not ausgewähltes_thema or not ausgewähltes_unterthema:
                    flash("Bitte Thema + Unterthema wählen.", 'warning')
                    return redirect(url_for('admin'))

                verarbeite_eintrag(eingabe_text, ausgewähltes_thema, ausgewähltes_unterthema, beschreibung_text)
                flash("Wissenseintrag gespeichert.", 'success')

            return redirect(url_for('admin'))

        if not themen_dict:
            flash("Bitte fügen Sie zuerst ein Thema hinzu.", 'warning')
            return redirect(url_for('admin'))
        return render_template('admin.html', themen_dict=themen_dict)

    except Exception as e:
        logging.exception("Fehler in admin-Funktion.")
        flash("Ein unerwarteter Fehler ist aufgetreten.", 'danger')
        return render_template('admin.html', themen_dict={})

@app.route('/edit', methods=['GET'])
@login_required
def edit():
    try:
        wissensbasis = download_wissensbasis()
        logging.debug("Wissensbasis geladen: %s", wissensbasis)
        return render_template('edit.html', wissensbasis=wissensbasis)
    except Exception as e:
        logging.exception("Fehler beim Laden der Bearbeitungsseite.")
        return "Interner Serverfehler", 500

@app.route('/get_unterthemen', methods=['POST'])
@login_required
def get_unterthemen():
    thema = request.form.get('thema')
    themen_dict = lade_themen()
    unterthemen_info = {}
    if thema in themen_dict:
        for key, value in themen_dict[thema].items():
            unterthemen_info[key] = {
                'title': value['title'],
                'beschreibung': value.get('beschreibung', '')
            }
    return jsonify({'unterthemen': unterthemen_info})

def verarbeite_eintrag(eingabe_text, ausgewähltes_thema, ausgewähltes_unterthema, beschreibung_text):
    messages = [
        {"role": "user", "content": "Du bist ein Experte, der aus Transkripten Wissen extrahiert..."},
        {"role": "user", "content": f"Extrahiere ohne Verluste:\n'''{eingabe_text}'''"}
    ]
    response = contact_openai(messages, model="o3-mini")
    if not response:
        flash("Fehler bei der Zusammenfassung.", 'danger')
        return
    eintrag = {
        "thema": ausgewähltes_thema,
        "unterthema": ausgewähltes_unterthema,
        "beschreibung": beschreibung_text,
        "inhalt": response
    }
    speichere_wissensbasis(eintrag)

@app.route('/update_entry', methods=['POST'])
@login_required
def update_entry():
    try:
        data = request.get_json()
        thema = data.get('thema', '').strip()
        unterthema = data.get('unterthema', '').strip()
        beschreibung = data.get('beschreibung', '').strip()
        inhalt = data.get('inhalt', '').strip()

        if not thema or not unterthema:
            return jsonify({'success': False, 'message': 'Ungültige Daten.'}), 400

        wissensbasis = download_wissensbasis()
        if thema in wissensbasis and unterthema in wissensbasis[thema]:
            wissensbasis[thema][unterthema]['beschreibung'] = beschreibung
            wissensbasis[thema][unterthema]['inhalt'] = inhalt.split('\n')
            upload_wissensbasis(wissensbasis)
            logging.debug(f"Eintrag '{unterthema}' in Thema '{thema}' aktualisiert.")
            return jsonify({'success': True}), 200
        else:
            wissensbasis.setdefault(thema, {})[unterthema] = {
                'beschreibung': beschreibung,
                'inhalt': inhalt.split('\n')
            }
            upload_wissensbasis(wissensbasis)
            logging.debug(f"Eintrag '{unterthema}' in Thema '{thema}' neu erstellt.")
            return jsonify({'success': True}), 200
    except Exception as e:
        logging.exception("Fehler beim Aktualisieren des Eintrags.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/move_entry', methods=['POST'])
@login_required
def move_entry():
    try:
        data = request.get_json()
        thema = data.get('thema')
        unterthema = data.get('unterthema')
        direction = data.get('direction')
        wissensbasis = download_wissensbasis()

        if thema not in wissensbasis or unterthema not in wissensbasis[thema]:
            return jsonify({'success': False, 'message': 'Eintrag nicht gefunden.'}), 404

        unterthemen = list(wissensbasis[thema].keys())
        idx = unterthemen.index(unterthema)

        if direction == 'up' and idx > 0:
            unterthemen[idx], unterthemen[idx-1] = unterthemen[idx-1], unterthemen[idx]
        elif direction == 'down' and idx < len(unterthemen) - 1:
            unterthemen[idx], unterthemen[idx+1] = unterthemen[idx+1], unterthemen[idx]
        else:
            return jsonify({'success': False, 'message': 'Verschieben nicht möglich.'}), 400

        wissensbasis[thema] = {k: wissensbasis[thema][k] for k in unterthemen}
        upload_wissensbasis(wissensbasis)
        logging.debug(f"Eintrag '{unterthema}' verschoben -> {direction}.")
        return jsonify({'success': True}), 200
    except Exception as e:
        logging.exception("Fehler beim Verschieben.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/delete_entry', methods=['POST'])
@login_required
def delete_entry():
    try:
        data = request.get_json()
        thema = data.get('thema')
        unterthema = data.get('unterthema')
        wissensbasis = download_wissensbasis()

        if thema in wissensbasis and unterthema in wissensbasis[thema]:
            del wissensbasis[thema][unterthema]
            upload_wissensbasis(wissensbasis)
            logging.debug(f"Eintrag '{unterthema}' gelöscht.")
            return jsonify({'success': True}), 200
        else:
            return jsonify({'success': False, 'message': 'Eintrag nicht gefunden.'}), 404
    except Exception as e:
        logging.exception("Fehler beim Löschen.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/sort_entries', methods=['POST'])
@login_required
def sort_entries():
    try:
        wissensbasis = download_wissensbasis()
        def sort_key(k):
            match = re.match(r'(\d+)([a-z]*)', k)
            if match:
                num = int(match.group(1))
                suf = match.group(2)
                return (num, suf)
            return (0, k)

        for thema, unterthemen in wissensbasis.items():
            sorted_keys = sorted(unterthemen.keys(), key=sort_key)
            wissensbasis[thema] = {k: wissensbasis[thema][k] for k in sorted_keys}
        upload_wissensbasis(wissensbasis)
        logging.debug("Wissensbasis sortiert.")
        return jsonify({'success': True}), 200
    except Exception as e:
        logging.exception("Fehler beim Sortieren.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/add_topic', methods=['POST'])
@login_required
def add_topic():
    try:
        data = request.get_json()
        add_type = data.get('type')

        if add_type == 'thema':
            neues_thema = data.get('thema', '').strip()
            if not neues_thema:
                return jsonify({'success': False, 'message': 'Thema darf nicht leer sein.'}), 400
            themen_dict = lade_themen()
            for thema in themen_dict.keys():
                if thema.startswith("Thema"):
                    existing_title = thema.split(':', 1)[1].strip()
                    if existing_title == neues_thema:
                        return jsonify({'success': False, 'message': 'Thema existiert bereits.'}), 400
            next_number = get_next_thema_number(themen_dict)
            thema_key = f'Thema {next_number}: {neues_thema}'
            themen_dict[thema_key] = {}
            aktualisiere_themen(themen_dict)
            flash(f'Hauptthema "{thema_key}" wurde erfolgreich hinzugefügt.', 'success')
            return jsonify({'success': True}), 200

        elif add_type == 'unterthema':
            parent_thema = data.get('parent_thema', '').strip()
            unterthema_nummer = data.get('unterthema_nummer', '').strip()
            unterthema_titel = data.get('unterthema_titel', '').strip()
            unterthema_beschreibung = data.get('unterthema_beschreibung', '').strip()
            if not parent_thema or not unterthema_nummer or not unterthema_titel:
                return jsonify({'success': False, 'message': 'Alle Felder sind erforderlich.'}), 400
            themen_dict = lade_themen()
            if parent_thema not in themen_dict:
                return jsonify({'success': False, 'message': 'Übergeordnetes Hauptthema existiert nicht.'}), 404
            if unterthema_nummer in themen_dict[parent_thema]:
                return jsonify({'success': False, 'message': 'Unterthema mit dieser Nummer existiert bereits.'}), 400
            themen_dict[parent_thema][unterthema_nummer] = {
                "title": unterthema_titel,
                "beschreibung": unterthema_beschreibung
            }
            aktualisiere_themen(themen_dict)
            flash(f'Unterthema "{unterthema_nummer}) {unterthema_titel}" hinzugefügt.', 'success')
            return jsonify({'success': True}), 200

        else:
            return jsonify({'success': False, 'message': 'Ungültiger Typ.'}), 400

    except Exception as e:
        logging.exception("Fehler beim Hinzufügen eines Themas.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/delete_topic', methods=['POST'])
@login_required
def delete_topic():
    try:
        data = request.get_json()
        thema = data.get('thema', '').strip()
        if not thema:
            return jsonify({'success': False, 'message': 'Thema darf nicht leer sein.'}), 400
        themen_dict = lade_themen()
        if thema not in themen_dict:
            return jsonify({'success': False, 'message': 'Thema existiert nicht.'}), 404
        del themen_dict[thema]
        aktualisiere_themen(themen_dict)
        flash(f'Thema "{thema}" wurde erfolgreich entfernt.', 'success')
        return jsonify({'success': True}), 200

    except Exception as e:
        logging.exception("Fehler beim Löschen eines Themas.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/upload_files', methods=['POST'])
@login_required
def upload_files():
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'message': 'Keine Dateien in der Anfrage.'}), 400
        files = request.files.getlist('files')
        if not files:
            return jsonify({'success': False, 'message': 'Keine Dateien ausgewählt.'}), 400

        if 'uploaded_files' not in session:
            session['uploaded_files'] = []
        files_status = []
        for file in files:
            filename = secure_filename(file.filename)
            if filename == '':
                continue
            file_id = str(uuid.uuid4())
            temp_filepath = os.path.join(app.config['UPLOAD_FOLDER'], file_id + '_' + filename)
            file.save(temp_filepath)
            file_status = {
                'id': file_id,
                'filename': filename,
                'status': 'Hochgeladen'
            }
            session['uploaded_files'].append(file_status)
            files_status.append(file_status)

        session.modified = True
        return jsonify({'success': True, 'files': files_status}), 200
    except Exception as e:
        logging.exception("Fehler beim Hochladen von Dateien.")
        return jsonify({'success': False, 'message': 'Interner Fehler.'}), 500

@app.route('/process_file_ai', methods=['POST'])
@login_required
def process_file_ai():
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        uploaded_files = session.get('uploaded_files', [])
        file_entry = next((f for f in uploaded_files if f['id'] == file_id), None)

        if not file_entry:
            return jsonify({'success': False, 'message': 'Datei nicht gefunden.'}), 404
        if file_entry['status'] != 'Hochgeladen':
            return jsonify({'success': False, 'message': 'Datei bereits verarbeitet.'}), 400

        temp_filepath = os.path.join(app.config['UPLOAD_FOLDER'], file_id + '_' + file_entry['filename'])
        if not os.path.exists(temp_filepath):
            return jsonify({'success': False, 'message': 'Temporäre Datei nicht gefunden.'}), 404

        ext = os.path.splitext(file_entry['filename'])[1].lower()
        extracted_text = ""
        if ext == '.txt':
            with open(temp_filepath, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
        elif ext == '.pdf':
            with open(temp_filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
        elif ext in ['.doc', '.docx']:
            doc = docx.Document(temp_filepath)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        else:
            return jsonify({'success': False, 'message': 'Unsupported file type.'}), 400

        if not extracted_text.strip():
            file_entry['status'] = 'Fehler: Kein Text extrahiert'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kein Text extrahiert.'}), 400

        themen_hierarchie = ''
        if os.path.exists(themen_datei):
            with open(themen_datei, 'r', encoding='utf-8') as f:
                themen_hierarchie = f.read()

        kategorisierung_messages = [
            {"role": "user", "content": "Du bist ein Assistent, der Texte in vorgegebene Themen..."},
            {"role": "user", "content": f"Hier die Themenhierarchie:\n\n{themen_hierarchie}\n\nText:\n{extracted_text}"}
        ]
        kategorisierung_response = contact_openai(kategorisierung_messages, model="o3-mini")
        if not kategorisierung_response:
            file_entry['status'] = 'Fehler: Kategorisierung fehlgeschlagen'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kategorisierung fehlgeschlagen.'}), 500

        try:
            kategorisierung_text = kategorisierung_response
            json_match = re.search(r'\[\s*{.*}\s*\]', kategorisierung_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(0)
                json_text = re.sub(r'(\d+[a-z]*)\)\)', r'\1)', json_text)
                kategorisierte_eintraege = json.loads(json_text)
            else:
                raise ValueError("Kein gültiger JSON-Inhalt gefunden")

            for eintrag_kat in kategorisierte_eintraege:
                thema = eintrag_kat.get('thema')
                unterthema = eintrag_kat.get('unterthema')
                beschreibung = eintrag_kat.get('beschreibung', '')
                inhalt = eintrag_kat.get('inhalt')
                if thema and unterthema and inhalt:
                    speichere_wissensbasis({
                        "thema": thema,
                        "unterthema": unterthema,
                        "beschreibung": beschreibung,
                        "inhalt": inhalt
                    })

            file_entry['status'] = 'Erfolgreich verarbeitet (KI)'
            session.modified = True
            os.remove(temp_filepath)
            return jsonify({'success': True, 'message': 'Datei erfolgreich verarbeitet.'}), 200
        except (ValueError, json.JSONDecodeError):
            file_entry['status'] = 'Fehler: Parsing-Fehler'
            session.modified = True
            return jsonify({'success': False, 'message': 'Parsing-Fehler bei der Kategorisierung.'}), 500

    except Exception as e:
        logging.exception("Fehler bei der automatischen Verarbeitung der Datei.")
        return jsonify({'success': False, 'message': 'Ein interner Fehler ist aufgetreten.'}), 500

@app.route('/process_file_manual', methods=['POST'])
@login_required
def process_file_manual():
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        selected_thema = data.get('thema')
        selected_unterthema = data.get('unterthema')
        beschreibung = data.get('beschreibung', '').strip()

        uploaded_files = session.get('uploaded_files', [])
        file_entry = next((f for f in uploaded_files if f['id'] == file_id), None)

        if not file_entry:
            return jsonify({'success': False, 'message': 'Datei nicht gefunden.'}), 404
        if file_entry['status'] != 'Hochgeladen':
            return jsonify({'success': False, 'message': 'Datei bereits verarbeitet.'}), 400

        temp_filepath = os.path.join(app.config['UPLOAD_FOLDER'], file_id + '_' + file_entry['filename'])
        if not os.path.exists(temp_filepath):
            return jsonify({'success': False, 'message': 'Temporäre Datei nicht gefunden.'}), 404

        ext = os.path.splitext(file_entry['filename'])[1].lower()
        extracted_text = ""
        if ext == '.txt':
            with open(temp_filepath, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
        elif ext == '.pdf':
            with open(temp_filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
        elif ext in ['.doc', '.docx']:
            doc = docx.Document(temp_filepath)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        else:
            return jsonify({'success': False, 'message': 'Unsupported file type.'}), 400

        if not extracted_text.strip():
            file_entry['status'] = 'Fehler: Kein Text extrahiert'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kein Text extrahiert.'}), 400

        speichere_wissensbasis({
            "thema": selected_thema,
            "unterthema": selected_unterthema,
            "beschreibung": beschreibung,
            "inhalt": extracted_text
        })

        file_entry['status'] = 'Erfolgreich verarbeitet (Manuell)'
        session.modified = True
        os.remove(temp_filepath)
        return jsonify({'success': True, 'message': 'Datei erfolgreich manuell verarbeitet.'}), 200

    except Exception as e:
        logging.exception("Fehler bei der manuellen Verarbeitung.")
        return jsonify({'success': False, 'message': 'Ein interner Fehler ist aufgetreten.'}), 500

###########################################
# App Start
###########################################
if __name__ == "__main__":
    app.run(debug=True)#