# routes/admin_file_processing.py

import os
import re
import json
import uuid
import logging
import traceback
from datetime import datetime
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, session, current_app
from werkzeug.utils import secure_filename
from functools import wraps
from app import (
    login_required, contact_openai, speichere_wissensbasis, themen_datei
)

try:
    from PyPDF2 import PdfReader  # For PDF parsing
except ImportError:
    # Placeholder if import fails
    PdfReader = None

try:
    import docx  # For DOCX parsing
except ImportError:
    # Placeholder if import fails
    docx = None


# Variables that need to be imported or configured
themen_datei = 'themen.txt'  # Placeholder path, should be imported or configured

# Blueprint definition
admin_file_processing_bp = Blueprint('admin_file_processing', __name__, url_prefix='/admin/files')


# Routes moved from app.py

@admin_file_processing_bp.route('/upload', methods=['POST'])
@login_required
def upload_files():
    """
    Handle file uploads for later processing.
    Files are temporarily stored and tracked in the session.
    """
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'message': 'Keine Dateien in der Anfrage.'}), 400
        
        files = request.files.getlist('files')
        if not files:
            return jsonify({'success': False, 'message': 'Keine Dateien ausgewählt.'}), 400

        # Initialize the uploaded_files session variable if it doesn't exist
        if 'uploaded_files' not in session:
            session['uploaded_files'] = []
            
        files_status = []
        
        # Process each uploaded file
        for file in files:
            filename = secure_filename(file.filename)
            if filename == '':
                continue
                
            # Generate a unique ID for the file
            file_id = str(uuid.uuid4())
            
            # Get the upload folder from app config or use a default
            upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
            
            # Ensure upload directory exists
            os.makedirs(upload_folder, exist_ok=True)
            
            # Save the file to the temporary location
            temp_filepath = os.path.join(upload_folder, file_id + '_' + filename)
            file.save(temp_filepath)
            
            # Create a status entry for this file
            file_status = {
                'id': file_id,
                'filename': filename,
                'status': 'Hochgeladen'
            }
            
            # Add to session and response
            session['uploaded_files'].append(file_status)
            files_status.append(file_status)

        # Mark session as modified to ensure it's saved
        session.modified = True
        
        return jsonify({
            'success': True, 
            'files': files_status,
            'message': f"{len(files_status)} Dateien erfolgreich hochgeladen."
        }), 200
        
    except Exception as e:
        logging.exception("Fehler beim Hochladen von Dateien.")
        return jsonify({'success': False, 'message': f'Interner Fehler: {str(e)}'}), 500


@admin_file_processing_bp.route('/process/ai', methods=['POST'])
@login_required
def process_file_ai():
    """
    Process a previously uploaded file using AI to categorize and extract content.
    The file will be automatically categorized into the appropriate topics.
    """
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'success': False, 'message': 'Keine Datei-ID angegeben.'}), 400
            
        # Find the file entry in the session
        uploaded_files = session.get('uploaded_files', [])
        file_entry = next((f for f in uploaded_files if f['id'] == file_id), None)

        if not file_entry:
            return jsonify({'success': False, 'message': 'Datei nicht gefunden.'}), 404
            
        if file_entry['status'] != 'Hochgeladen':
            return jsonify({'success': False, 'message': 'Datei bereits verarbeitet.'}), 400

        # Get path to the temporary file
        upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
        temp_filepath = os.path.join(upload_folder, file_id + '_' + file_entry['filename'])
        
        if not os.path.exists(temp_filepath):
            return jsonify({'success': False, 'message': 'Temporäre Datei nicht gefunden.'}), 404

        # Extract text based on file extension
        ext = os.path.splitext(file_entry['filename'])[1].lower()
        extracted_text = ""
        
        if ext == '.txt':
            with open(temp_filepath, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
                
        elif ext == '.pdf':
            if not PdfReader:
                return jsonify({'success': False, 'message': 'PDF-Unterstützung nicht verfügbar.'}), 500
            
            with open(temp_filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
                    
        elif ext in ['.doc', '.docx']:
            if not docx:
                return jsonify({'success': False, 'message': 'DOCX-Unterstützung nicht verfügbar.'}), 500
                
            doc = docx.Document(temp_filepath)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
                
        else:
            return jsonify({
                'success': False, 
                'message': f'Dateityp nicht unterstützt: {ext}. Unterstützte Typen: .txt, .pdf, .doc, .docx'
            }), 400

        # Check if any text was extracted
        if not extracted_text.strip():
            file_entry['status'] = 'Fehler: Kein Text extrahiert'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kein Text aus der Datei extrahiert.'}), 400

        # Load theme hierarchy for categorization
        themen_hierarchie = ''
        if os.path.exists(themen_datei):
            with open(themen_datei, 'r', encoding='utf-8') as f:
                themen_hierarchie = f.read()

        # Use AI to categorize the extracted text
        kategorisierung_messages = [
            {"role": "user", "content": "Du bist ein Assistent, der Texte in vorgegebene Themen..."},
            {"role": "user", "content": f"Hier die Themenhierarchie:\n\n{themen_hierarchie}\n\nText:\n{extracted_text}"}
        ]
        
        kategorisierung_response = contact_openai(kategorisierung_messages, model="gpt-4o")
        
        if not kategorisierung_response:
            file_entry['status'] = 'Fehler: Kategorisierung fehlgeschlagen'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kategorisierung fehlgeschlagen.'}), 500

        try:
            # Extract JSON from the AI response
            kategorisierung_text = kategorisierung_response
            json_match = re.search(r'\[\s*{.*}\s*\]', kategorisierung_text, re.DOTALL)
            
            if json_match:
                json_text = json_match.group(0)
                # Fix potential formatting issues in AI response
                json_text = re.sub(r'(\d+[a-z]*)\)\)', r'\1)', json_text)
                kategorisierte_eintraege = json.loads(json_text)
            else:
                raise ValueError("Kein gültiger JSON-Inhalt in der KI-Antwort gefunden")

            # Process each categorized entry
            for eintrag_kat in kategorisierte_eintraege:
                thema = eintrag_kat.get('thema')
                unterthema = eintrag_kat.get('unterthema')
                beschreibung = eintrag_kat.get('beschreibung', '')
                inhalt = eintrag_kat.get('inhalt')
                
                if thema and unterthema and inhalt:
                    # Save to knowledge base
                    speichere_wissensbasis({
                        "thema": thema,
                        "unterthema": unterthema,
                        "beschreibung": beschreibung,
                        "inhalt": inhalt
                    })

            # Update status and clean up
            file_entry['status'] = 'Erfolgreich verarbeitet (KI)'
            session.modified = True
            
            # Delete temporary file
            os.remove(temp_filepath)
            
            return jsonify({
                'success': True, 
                'message': f'Datei erfolgreich verarbeitet und in Wissensbasis integriert.'
            }), 200
            
        except (ValueError, json.JSONDecodeError) as e:
            file_entry['status'] = 'Fehler: Parsing-Fehler'
            session.modified = True
            return jsonify({
                'success': False, 
                'message': f'Parsing-Fehler bei der Kategorisierung: {str(e)}'
            }), 500

    except Exception as e:
        logging.exception("Fehler bei der automatischen Verarbeitung der Datei.")
        return jsonify({
            'success': False, 
            'message': f'Ein interner Fehler ist aufgetreten: {str(e)}'
        }), 500


@admin_file_processing_bp.route('/process/manual', methods=['POST'])
@login_required
def process_file_manual():
    """
    Process a previously uploaded file by manually assigning it to a topic/subtopic.
    This skips the AI categorization step.
    """
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        selected_thema = data.get('thema')
        selected_unterthema = data.get('unterthema')
        beschreibung = data.get('beschreibung', '').strip()

        # Validate inputs
        if not file_id:
            return jsonify({'success': False, 'message': 'Keine Datei-ID angegeben.'}), 400
            
        if not selected_thema or not selected_unterthema:
            return jsonify({'success': False, 'message': 'Thema und Unterthema müssen angegeben werden.'}), 400

        # Find the file entry in the session
        uploaded_files = session.get('uploaded_files', [])
        file_entry = next((f for f in uploaded_files if f['id'] == file_id), None)

        if not file_entry:
            return jsonify({'success': False, 'message': 'Datei nicht gefunden.'}), 404
            
        if file_entry['status'] != 'Hochgeladen':
            return jsonify({'success': False, 'message': 'Datei bereits verarbeitet.'}), 400

        # Get path to the temporary file
        upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
        temp_filepath = os.path.join(upload_folder, file_id + '_' + file_entry['filename'])
        
        if not os.path.exists(temp_filepath):
            return jsonify({'success': False, 'message': 'Temporäre Datei nicht gefunden.'}), 404

        # Extract text based on file extension
        ext = os.path.splitext(file_entry['filename'])[1].lower()
        extracted_text = ""
        
        if ext == '.txt':
            with open(temp_filepath, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
                
        elif ext == '.pdf':
            if not PdfReader:
                return jsonify({'success': False, 'message': 'PDF-Unterstützung nicht verfügbar.'}), 500
                
            with open(temp_filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
                    
        elif ext in ['.doc', '.docx']:
            if not docx:
                return jsonify({'success': False, 'message': 'DOCX-Unterstützung nicht verfügbar.'}), 500
                
            doc = docx.Document(temp_filepath)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
                
        else:
            return jsonify({
                'success': False, 
                'message': f'Dateityp nicht unterstützt: {ext}. Unterstützte Typen: .txt, .pdf, .doc, .docx'
            }), 400

        # Check if any text was extracted
        if not extracted_text.strip():
            file_entry['status'] = 'Fehler: Kein Text extrahiert'
            session.modified = True
            return jsonify({'success': False, 'message': 'Kein Text aus der Datei extrahiert.'}), 400

        # Save to knowledge base with manual categorization
        speichere_wissensbasis({
            "thema": selected_thema,
            "unterthema": selected_unterthema,
            "beschreibung": beschreibung,
            "inhalt": extracted_text
        })

        # Update status and clean up
        file_entry['status'] = 'Erfolgreich verarbeitet (Manuell)'
        session.modified = True
        
        # Delete temporary file
        os.remove(temp_filepath)
        
        return jsonify({
            'success': True, 
            'message': f'Datei wurde erfolgreich verarbeitet und in {selected_thema}/{selected_unterthema} gespeichert.'
        }), 200

    except Exception as e:
        logging.exception("Fehler bei der manuellen Verarbeitung.")
        return jsonify({
            'success': False, 
            'message': f'Ein interner Fehler ist aufgetreten: {str(e)}'
        }), 500
